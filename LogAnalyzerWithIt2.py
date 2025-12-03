#=====================================================
# LogAnalyzerNew.py
# Jake Chinchar
# Purpose: Analyze Communication Manager logs, Wireshark PCAP files, and Packetswitch HTML reports to correlate message flows.
# Output: Generates a formatted Excel report summarizing findings
# Date last edited: 12/3/2025
#=====================================================

# Wireshark Filter Example: 
# ((atcsl3.srce_addr contains 71:a3:a7:8a:36) || (atcsl3.dest_addr contains 71:a3:a7:8a:36)) && !(_ws.col.protocol == "ATCS/TCP")

# For Git
# cd "folder path"
# git status
# git add LogAnalyzerWithIt2.py
# git commit -m "Describe changes"
# git push origin main

# Required External Packages:
# pip install scapy pandas openpyxl beautifulsoup4 python-docx

# Built-in modules: tkinter, datetime, re, os, subprocess, json

# CODE IMPROVEMENTS TO IMPLEMENT IN THE FUTURE:
# 1. Add date filtering
# 2. Find addresses in wireshark to choose from (the user should not have to enter the address, they should select)
# 3. Incoorporate an incorrect input file notification
# 4. Have a progress bar
# 5. Check to see if the ATCS address in the comms manager log matches the one entered (if a comms manager file is provided)
# 6. Eliminate the need for a wireshark file
# 7. Prevent errors for when user enteres an output file (/ and ? cause errors as of now)

# Import Libraries
import tkinter as tk                            # For GUI
from tkinter import filedialog, messagebox      # filedialog for opening/saving files, messagebox for error alerts and success notifs
from scapy.all import rdpcap, Raw               # Read packets from pcap file
from scapy.layers.inet import TCP               # To identify TCP packets
from datetime import datetime, timedelta        # For dates and times
import re                                       # Pattern matching & parsing
import os                                       # Operating system functions
import pandas as pd                             # For dataframes
from openpyxl import load_workbook              # To write to and modify excel files
from openpyxl.styles import Alignment, PatternFill           # For excel alignment
from openpyxl.utils import get_column_letter    # Converts column index to Excel column letters
from bs4 import BeautifulSoup                   # To read HTML files
from docx import Document                       # To read .docx
import subprocess                               # Used to run external commands like 'tshark' for extracting Wireshark data    
import json                                     # Used to parse JSON output from tshark and handle structured data
from openpyxl.utils import get_column_letter    # Converts numeric column index to Excel letters (e.g., 1 -> 'A') for formatting
from openpyxl.styles import Alignment, PatternFill           # Used to set cell alignment in Excel (center, left, wrap text)

            
# Expression to extract time tags from log lines (e.g., 12:34:56.78)
time_pattern = re.compile(r'\b\d{2}:\d{2}:\d{2}\.\d{2}\b')

# Converts a time string to a datetime object
def parse_time_only(ts_str):
    # Add a .0 if the user does not add a decimal to the input start or end time
    if '.' not in ts_str:
        ts_str += '.0'

    # Hours:Minutes:Seconds.Decimal
    return datetime.strptime(ts_str, '%H:%M:%S.%f')

# Formats a timedelta object into a readable string (HH:MM:SS.ms)
def format_timedelta(td):
    total_seconds = int(td.total_seconds())
    milliseconds = int(td.microseconds / 10000)
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours:02}:{minutes:02}:{seconds:02}.{milliseconds:02}"

# Checks if a packet contains the target address and is not TCP
# FUTURE CHANGE: TCP will be valid, will instead specify in the output file if it is TCP or not, so that the user can filter themselves
def is_valid_packet(pkt, data_bytes, target_address):
    return target_address in data_bytes and not TCP in pkt

# Searches for a matching packet in the pcap file
# Finds the matching packet by looking for the first entry with the same message type and message number within 8 minutes of either direction of the comms manager time tag
def find_pcap_time(search_bytes, log_dt, packets, target_address, pcap_file_path, msg_type_raw):
    for pkt in packets:
        if Raw in pkt and target_address in pkt[Raw].load:
            data_bytes = pkt[Raw].load
            if search_bytes in data_bytes and target_address in data_bytes and not TCP in pkt:
                pcap_time = float(pkt.time)
                pcap_dt = datetime.fromtimestamp(pcap_time)
                pcap_ts_str = pcap_dt.strftime('%H:%M:%S.%f')[:-3]

                try:
                    pcap_time_only = datetime.strptime(pcap_ts_str, '%H:%M:%S.%f')
                    
                    time_diff = pcap_time_only - log_dt
                    minutes_diff = time_diff.total_seconds() / 60
                    # Valid minutes difference allowed (wireshark time tag can be -8 minutes before or 8 minutes
                    
                    # Filter to ensure that the entry is found is within a realistic time frame
                    if pcap_time_only < (log_dt - timedelta(minutes=8)):
                        continue

                    # Valid time range of messages (comms manager time tags are not accurate and this may need to be changed)
                    if -8 <= minutes_diff <= 8:
                        fallback_hex = extract_ws_hex_data('', '', '', fallback_hex='')  # placeholder removed
                        ws_hex_data = extract_ws_hex_data(pcap_file_path, pcap_ts_str, msg_type_raw,
                                                          fallback_hex=extract_ws_hex_data('', '', '', fallback_hex=''))
                        # Actually, fallback_hex should be old raw slice:
                        fallback_hex = ' '.join([data_bytes.hex()[i:i+2] for i in range(0, len(data_bytes.hex()), 2)])
                        ws_hex_data = extract_ws_hex_data(pcap_file_path, pcap_ts_str, msg_type_raw, fallback_hex)
                        return pcap_ts_str, format_timedelta(time_diff), True, ws_hex_data
                    else:
                        # The wireshark message found does not correspond to the communication manager message with the same msg number and type
                        # Messages with "Found but overflow" are not going to be printed to the excel file
                        # There is a finite number of message numbers, so this states that there was wrapping of the message number 
                            return pcap_ts_str, "Found but overflow", False, ''
                except:
                    continue

    # If no messages or overflow were found, then the message is "LOST", which means that the communication manager log has this message, but the wireshark did not
    return None, "LOST", False, ''
        

# Checks if the output Excel file is currently open/locked
# This script does not write over output files that are open. The user can close the file or wire a different name for the output file
def is_file_open(filepath):
    try:
        os.rename(filepath, filepath)
        return False
    except:
        return True



def extract_ws_hex_data(pcap_path: str, ws_time_tag: str, msg_type_raw: str, fallback_hex: str = '') -> str:
    """
    Extracts a specific hex payload from a Wireshark PCAP using tshark JSON output,
    matching exactly on a Wireshark-formatted timestamp (HH:MM:SS.mmm).
    
    Parameters
    ----------
    pcap_path : str
        Full path to the .pcap file.
    ws_time_tag : str
        Target Wireshark time tag in 'HH:MM:SS.mmm' format (milliseconds).
        Example: '12:34:56.789'
    msg_type_raw : str
        Raw message type from the CM/WS context. Expected values:
        - '12 8B' -> we look for 'indication_bits' in the tshark JSON
        - '12 01' -> we look for 'control_bits' in the tshark JSON
    fallback_hex : str, optional
        Hex string to return if we cannot find (or parse) the desired payload.
        This ensures the caller still gets something usable.

    Returns
    -------
    str
        Space-separated uppercase hex pairs (e.g., 'DE AD BE EF').
        Returns `fallback_hex` if no matching frame/payload is found.
    """

    # Map the raw message type to the corresponding label name that appears in
    # tshark's JSON "showname" fields. If type isn't recognized, we bail early.
    if msg_type_raw == '12 8B':
        wanted_label = 'indication_bits'
    elif msg_type_raw == '12 01':
        wanted_label = 'control_bits'
    else:
        return ''  # Unknown msg_type_raw: no defined label to search for

    try:
        # Phase 1: Narrow the search window to the matching second
        # We only filter by the 'HH:MM:SS' part to reduce I/O and JSON size.
        # Later, we verify exact millisecond match.
        second_str = ws_time_tag.split('.')[0]  # 'HH:MM:SS'
        display_filter = f'frame.time contains "{second_str}"'

        # Run tshark to produce JSON for frames in that second.
        # -r <pcap> : read file
        # -Y <filter> : display filter (not capture filter)
        # -T json : output as JSON
        proc = subprocess.run(
            ["tshark", "-r", pcap_path, "-Y", display_filter, "-T", "json"],
            capture_output=True, text=True, check=True  # check=True raises CalledProcessError on non-zero exit
        )
        frames = json.loads(proc.stdout)  # List[dict] of frame objects

        # Local import (already available globally), used to format epoch to HH:MM:SS.mmm
        from datetime import datetime

        # Phase 2: Find the single frame that exactly matches 'ws_time_tag' 
        # Read each frame, get epoch time → convert to HH:MM:SS.mmm, compare.
        for fr in frames:
            layers = fr.get("_source", {}).get("layers", {})  # Safe access to nested JSON
            t_epoch_list = layers.get("frame", {}).get("frame.time_epoch", [])
            if not t_epoch_list:
                continue  # No epoch in this frame; skip

            try:
                # Convert epoch to float → datetime → format 'HH:MM:SS.mmm'
                t_epoch = float(t_epoch_list[0])
                candidate = datetime.fromtimestamp(t_epoch).strftime('%H:%M:%S.%f')[:-3]
            except:
                # If any parsing error happens, ignore this frame
                continue

            # Only proceed for the exact millisecond match (same 'HH:MM:SS.mmm')
            if candidate != ws_time_tag:
                continue

            # Phase 3: Walk the JSON tree to locate the value tied to `wanted_label` 
            # tshark JSON can be deeply nested; we recurse through dicts/lists.
            # We look specifically for a field with 'showname' that starts with '<wanted_label>: '.
            def find_val(node):
                # Recursive search across dicts/lists for a matching 'showname'
                if isinstance(node, dict):
                    if 'showname' in node and isinstance(node['showname'], str):
                        prefix = f"{wanted_label}: "  # e.g., 'indication_bits: ' or 'control_bits: '
                        if node['showname'].startswith(prefix):
                            # Extract everything after '<label>: '
                            return node['showname'][len(prefix):].strip()
                    # Recurse into nested values
                    for v in node.values():
                        val = find_val(v)
                        if val:
                            return val
                elif isinstance(node, list):
                    # Recurse into each item of the list
                    for item in node:
                        val = find_val(item)
                        if val:
                            return val
                return None  # Not found at this branch

            val = find_val(layers)
            if val:
                # Phase 4: Normalize to hex pairs 
                # Remove any non-hex characters, then group into uppercased byte pairs.
                raw = re.sub(r'[^0-9A-Fa-f]', '', val)
                return ' '.join(raw[i:i+2].upper() for i in range(0, len(raw), 2))

    except:
        # Any error (tshark not present, invalid JSON, CalledProcessError, etc.)
        # Do not crash—return fallback hex instead.
        pass

    # If we got here, we didn't find a precise match or parsing failed—use fallback
    return fallback_hex



# Helper function to process Wireshark hex data
# It searches for '02 02 12 8B' or '02 02 12 01', removes everything before and including that sequence, then removes the next four hex pairs.
# The data is always in this format

# === Helper functions for message type normalization ===
def _normalize_group_from_raw(msg_type_raw: str) -> str:
    s = (msg_type_raw or '').upper().strip()
    if s in ('12 8B', '78 8B'):
        return 'IND'
    if s in ('12 01', '12 48'):
        return 'CTL_RECALL'
    if s in ('08 82', '08 83'):
        return 'GROUP_0882_0883'
    return s

def _group_from_msg_type_cell(cell_val: str) -> str:
    val = (cell_val or '').strip()
    if '(' in val and ')' in val:
        raw = val.split('(')[-1].strip(')')
    else:
        raw = val
    return _normalize_group_from_raw(raw)

def process_ws_hex_data(ws_hex_data: str, msg_type_raw: str) -> str:
    if not ws_hex_data:
        return ws_hex_data
    parts = ws_hex_data.strip().split()

    # Find index of sequence
    idx = None
    for i in range(len(parts) - 3):
        if parts[i:i+4] == ['02', '02', '12', '8b'] or parts[i:i+4] == ['02', '02', '12', '01']:
            idx = i + 4  # position after the sequence
            break
    if idx is None:
        return ws_hex_data  # no change if sequence not found
    # Remove everything before and including sequence
    trimmed = parts[idx:]
    # Remove next four pairs if available
    trimmed = trimmed[4:] if len(trimmed) > 4 else []

    # If the message type is 12 8b, then the last hexadecimal pair needs to be removed
    if msg_type_raw == '12 8B' and len(trimmed) > 0:
        trimmed = trimmed[:-1]

    return ' '.join(trimmed)




def extract_hex_data(lines, start_index, msg_type_raw, line_type):
    """
    1. After finding '02 02 12 8B' or '02 02 12 01', remove first 2 hex pairs.
    2. Take the next (third) hex pair after message type, convert to binary.
    3. Remove first 2 bits from that binary value.
    4. Use remaining bits as length (in hex pairs).
    5. Trim collected hex pairs to match this length.
    """
    
    collected = []
    found_sequence = False
    sequence_patterns = [['02', '02', '12', '8B'], ['02', '02', '12', '01']]

    i = start_index
    while i < len(lines):
        tokens = lines[i].strip().split()
        joined_line = ' '.join(tokens).upper()
        if ' TX' in joined_line or ' RX' in joined_line or 'TX:' in joined_line or 'RX:' in joined_line:
            break

        if not found_sequence:
            for j in range(len(tokens) - 3):
                if tokens[j:j+4] in sequence_patterns:
                    found_sequence = True
                    collected.extend([t for t in tokens[j+4:] if re.fullmatch(r'[0-9A-Fa-f]{2}', t)])
                    break
        else:
            collected.extend([t for t in tokens if re.fullmatch(r'[0-9A-Fa-f]{2}', t)])

        i += 1

    if not found_sequence or len(collected) < 3:
        return ''

    # Remove first 2 pairs
    collected = collected[2:]

    # Get third pair (now first after removal) for length calculation
    length_hex = collected[0]
    length_bin = bin(int(length_hex, 16))[2:].zfill(8)  # Convert to 8-bit binary
    length_bin_trimmed = length_bin[2:]  # Remove first 2 bits
    length_val = int(length_bin_trimmed, 2)  # Convert back to int

    # Remove the length byte itself from data
    collected = collected[2:]

    # Trim to required length
    if len(collected) > length_val:
        collected = collected[:length_val]

    return ' '.join(collected)



#=======================
# IXL code section start
#=======================

# Future: move this section to its own file for better clarity

# XL parsing & matching (RULE-BASED COMPONENTS, NEWLINES)
# - Groups CTL/IND rows (per-row bit reversal).
# - Applies IXL start/end time filter.
# - Matches to Wireshark by same msg type + exact payload + |Δt| <= 12 minutes.
# - COMPONENT RULE:
#     For any line containing " -- ", take the substring AFTER " -- ".
#     If it does NOT start with IND/CTL/Execute (case-insensitive), capture it as a component line
#     and attach it to the most recently completed Control or Ind message.
#     Place each captured component on its own line within the same Excel cell.

IXL_LABEL_TO_WS_TYPE = {
    'CTL': 'Control (12 01)',
    'IND': 'Ind (12 8B)',
}


ixl_line_re = re.compile(
    r'^(?:\w{3}\s+)?'                # Optional weekday (e.g., Wed)
    r'(?P<md>\d{2}-\d{2}(?:-\d{4})?)\s+'  # MM-DD or MM-DD-YYYY
    r'(?P<hms>\d{2}:\d{2}:\d{2}\.\d{2}).*?\b'
    r'(?P<label>CTL|IND)\(\d{1,3}-\d{1,3}\):\s+(?P<bits>[01]{8})'
)

# ^ Regex to parse IXL log lines into structured components:
# Example line: "11-20 14:32:07.45 ... CTL(12-34): 10101100"
# Captures:
#   md    -> Month-Day (e.g., "11-20")
#   hms   -> Time in HH:MM:SS.xx format (e.g., "14:32:07.45")
#   label -> Message type: "CTL" (Control) or "IND" (Indication)
#   bits  -> 8-bit binary string (e.g., "10101100")
# Pattern details:
#   (?P<md>\d{2}-\d{2})                 : Two digits, dash, two digits (date)
#   \s+                                 : Whitespace separator
#   (?P<hms>\d{2}:\d{2}:\d{2}\.\d{2})   : Time with fractional seconds
#   .*?                                 : Match for intermediate text
#   \b(?P<label>CTL|IND)                : Word boundary, then CTL or IND
#   \(\d{1,3}-\d{1,3}\):                : Message number range in parentheses

ixl_end_execute_re = re.compile(r'Execute issued', re.IGNORECASE)

MAX_IXL_WS_DIFF_MINUTES = 8
# Max time difference allowed between IXL and Wireshark entries

def parse_time_flexible(ts_str: str) -> datetime:
    if not ts_str:
        raise ValueError("Empty time string")
    if '.' in ts_str:
        hms, frac = ts_str.split('.', 1)
        if len(frac) == 2:
            ts_norm = f"{hms}.{frac}0"
        elif len(frac) >= 3:
            ts_norm = f"{hms}.{frac[:3]}"
        else:
            ts_norm = f"{hms}.{frac.ljust(3, '0')}"
    else:
        ts_norm = ts_str + ".000"
    return datetime.strptime(ts_norm, "%H:%M:%S.%f")

def reverse_bits_to_hex(bits_8: str) -> str:
    rev = bits_8[::-1]
    return f"{int(rev, 2):02X}"

def bits_list_to_hex(bits_list: list[str]) -> str:
    return ' '.join(reverse_bits_to_hex(b) for b in bits_list)

def normalize_hex_no_spaces(s: str) -> str:
    return (s or '').replace(' ', '').upper()

def minutes_abs(dt_a: datetime, dt_b: datetime) -> float:
    a = dt_a.time(); b = dt_b.time()
    to_secs = lambda t: t.hour*3600 + t.minute*60 + t.second + t.microsecond/1_000_000
    return abs(to_secs(a) - to_secs(b)) / 60.0


def _extract_component_from_line(line: str) -> str | None:
    """
    Extract component text from either:
      1) Legacy lines that include " -- "
      2) Timestamped lines with optional weekday/year (e.g. "Wed 11-12-2025 04:00:03.76  GEO_BOX_1W: ...")

    Lines starting with IND/CTL or "Execute issued" are ignored.
    """
    # Try legacy " -- " anchor first
    anchor = line.find(' -- ')
    if anchor != -1:
        after = line[anchor + len(' -- '):].strip()
    else:
        # Fallback: strip leading timestamp (optional weekday + optional year)
        m = re.match(
            r'^(?:\w{3}\s+)?'                  # optional weekday, e.g., 'Wed '
            r'\d{2}-\d{2}(?:-\d{4})?\s+'       # MM-DD or MM-DD-YYYY
            r'\d{2}:\d{2}:\d{2}\.\d{2}\s+'     # HH:MM:SS.ff
            r'(?P<rest>.+)$'                   # the remainder is candidate component text
        , line)
        if not m:
            return None
        after = m.group('rest').strip()

    if not after:
        return None

    # Ignore non-component starters
    head = after[:16].lower()
    if head.startswith('ind(') or head.startswith('ctl(') or head.startswith('execute'):
        return None

    return after



def _flush_ixl_group(acc_label, acc_bits, acc_time, out_msgs,
                     start_dt: datetime | None, end_dt: datetime | None,
                     acc_components: list[str] | None = None) -> int | None:
    """Append flushed message and return its index, or None if filtered out."""
    if not acc_label or not acc_bits:
        return None

    ws_type = IXL_LABEL_TO_WS_TYPE.get(acc_label.upper())
    if not ws_type:
        return None

    include = True
    if start_dt or end_dt:
        try:
            t = parse_time_flexible(acc_time)
            if start_dt and t.time() < start_dt.time():
                include = False
            if end_dt and t.time() > end_dt.time():
                include = False
        except Exception:
            include = False

    if not include:
        return None

    data_hex = bits_list_to_hex(acc_bits)
    comp_lines = acc_components or []  # NEW: carry pending component lines

    out_msgs.append({
        "time_tag": acc_time or '',
        "msg_type": ws_type,
        "data_hex": data_hex,
        "_data_norm": normalize_hex_no_spaces(data_hex),
        "component_lines": comp_lines,  # NEW: attach here
    })
    return len(out_msgs) - 1



def parse_ixl_file(ixl_file_path: str,
                   start_dt: datetime | None = None,
                   end_dt: datetime | None = None) -> list[dict]:
    """
    Parse IXL into grouped CTL/IND messages (per-row bit reversal).

    Merge rule:
      - If there is a raw data line (IND/CTL) in the next 2 lines whose
        time tag is within ±0.1 seconds of the first raw data line's time tag
        (and label matches), treat it as part of the same raw data block.

    Component rule:
      - While accumulating a block, if a non-data line's time tag is within
        ±0.1 seconds of the block's start time, attach it as a component
        to the same block (do NOT flush).
      - After a block is flushed, subsequent component lines are still
        attached via the existing "awaiting_components" behavior until
        the next CTL/IND row starts.
    """
    if not ixl_file_path or not os.path.isfile(ixl_file_path):
        return []

    def _within_point_one_seconds(t1_str: str, t2_str: str) -> bool:
        """Return True if |t1 - t2| <= 0.1 seconds."""
        try:
            t1 = parse_time_flexible(t1_str)
            t2 = parse_time_flexible(t2_str)
            return abs((t2 - t1).total_seconds()) <= 0.1
        except Exception:
            return False

    # Use the global time_pattern to read time from any line (including components)
    def _time_from_any_line(s: str) -> str | None:
        m = time_pattern.search(s)
        return m.group(0) if m else None

    messages: list[dict] = []
    acc_label: str | None = None
    acc_bits: list[str] = []
    acc_time: str | None = None
    acc_components: list[str] = []  # NEW: hold component lines until flush
    awaiting_components: bool = False
    last_msg_idx: int | None = None

    # index-based processing to allow safe lookahead
    with open(ixl_file_path, 'r', encoding='utf-8', errors='ignore') as f:
        lines = f.readlines()

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip('\n')

        # 1) CTL/IND data row?
        m = ixl_line_re.search(line)
        if m:
            # stop post-flush component capture; we're inside a data block
            awaiting_components = False
            last_msg_idx = None

            label = m.group('label').upper()
            bits  = m.group('bits')
            ttag  = m.group('hms')  # e.g., '04:16:06.00'

            if acc_label is None:
                # start a new accumulation block
                acc_label = label
                acc_time  = ttag
                acc_bits.append(bits)
                acc_components = []

                # merge up to next two IND/CTL rows within ±0.1s and same label
                j = i
                lookahead_count = 0
                while lookahead_count < 2 and (j + 1) < n:
                    nxt = lines[j + 1].rstrip('\n')
                    nm  = ixl_line_re.search(nxt)
                    if nm:
                        nxt_label = nm.group('label').upper()
                        nxt_time  = nm.group('hms')
                        if nxt_label == label and _within_point_one_seconds(ttag, nxt_time):
                            acc_bits.append(nm.group('bits'))
                            j += 1
                            lookahead_count += 1
                            continue
                    # if next line is a component with same time window, capture it but don't advance lookahead
                    comp = _extract_component_from_line(nxt)
                    if comp:
                        comp_time = _time_from_any_line(nxt)
                        if comp_time and _within_point_one_seconds(ttag, comp_time):
                            acc_components.append(comp)
                            j += 1
                            # do NOT increment lookahead_count; only IND/CTL rows count
                            continue
                    break

                i = j + 1
                continue

            else:
                # already accumulating; decide merge vs flush
                if label == acc_label and _within_point_one_seconds(acc_time, ttag):
                    # continuation of same block
                    j = i
                    acc_bits.append(bits)

                    lookahead_count = 0
                    while lookahead_count < 2 and (j + 1) < n:
                        nxt = lines[j + 1].rstrip('\n')
                        nm  = ixl_line_re.search(nxt)
                        if nm:
                            nxt_label = nm.group('label').upper()
                            nxt_time  = nm.group('hms')
                            if nxt_label == acc_label and _within_point_one_seconds(acc_time, nxt_time):
                                acc_bits.append(nm.group('bits'))
                                j += 1
                                lookahead_count += 1
                                continue
                        # same-time component while block is open -> collect, don't flush
                        comp = _extract_component_from_line(nxt)
                        if comp:
                            comp_time = _time_from_any_line(nxt)
                            if comp_time and _within_point_one_seconds(acc_time, comp_time):
                                acc_components.append(comp)
                                j += 1
                                # components do not consume lookahead quota
                                continue
                        break

                    i = j + 1
                    continue
                else:
                    # different label or outside 0.1s -> flush previous block
                    idx = _flush_ixl_group(acc_label, acc_bits, acc_time,
                                           messages, start_dt, end_dt,
                                           acc_components)  # NEW
                    if idx is not None and (
                        messages[idx]['msg_type'] == 'Control (12 01)' or
                        messages[idx]['msg_type'] == 'Ind (12 8B)'
                    ):
                        awaiting_components = True
                        last_msg_idx = idx
                    else:
                        awaiting_components = False
                        last_msg_idx = None

                    # start a new block for current row
                    acc_label, acc_bits, acc_time = label, [], ttag
                    acc_components = []
                    j = i
                    acc_bits.append(bits)

                    lookahead_count = 0
                    while lookahead_count < 2 and (j + 1) < n:
                        nxt = lines[j + 1].rstrip('\n')
                        nm  = ixl_line_re.search(nxt)
                        if nm:
                            nxt_label = nm.group('label').upper()
                            nxt_time  = nm.group('hms')
                            if nxt_label == acc_label and _within_point_one_seconds(acc_time, nxt_time):
                                acc_bits.append(nm.group('bits'))
                                j += 1
                                lookahead_count += 1
                                continue
                        comp = _extract_component_from_line(nxt)
                        if comp:
                            comp_time = _time_from_any_line(nxt)
                            if comp_time and _within_point_one_seconds(acc_time, comp_time):
                                acc_components.append(comp)
                                j += 1
                                continue
                        break

                    i = j + 1
                    continue

        # 2) "Execute issued"? Flush block but don't treat as a component line
        if ixl_end_execute_re.search(line):
            idx = _flush_ixl_group(acc_label, acc_bits, acc_time,
                                   messages, start_dt, end_dt,
                                   acc_components)  # NEW
            if idx is not None and (
                messages[idx]['msg_type'] == 'Control (12 01)' or
                messages[idx]['msg_type'] == 'Ind (12 8B)'
            ):
                awaiting_components = True
                last_msg_idx = idx
            else:
                awaiting_components = False
                last_msg_idx = None
            acc_label, acc_bits, acc_time = None, [], None
            acc_components = []
            i += 1
            continue

        # 3) Non-data line while accumulating:
        #    If its time is within ±0.1s of acc_time -> attach to current block (do NOT flush).
        #    Else -> flush and treat via post-flush component rule.
        if acc_bits:
            comp = _extract_component_from_line(line)
            if comp:
                comp_time = _time_from_any_line(line)
                if comp_time and acc_time and _within_point_one_seconds(acc_time, comp_time):
                    # same-time component -> keep block open, collect
                    acc_components.append(comp)
                    i += 1
                    continue

            # not a same-time component -> flush block first
            idx = _flush_ixl_group(acc_label, acc_bits, acc_time,
                                   messages, start_dt, end_dt,
                                   acc_components)  # NEW
            if idx is not None and (
                messages[idx]['msg_type'] == 'Control (12 01)' or
                messages[idx]['msg_type'] == 'Ind (12 8B)'
            ):
                awaiting_components = True
                last_msg_idx = idx
            else:
                awaiting_components = False
                last_msg_idx = None
            acc_label, acc_bits, acc_time = None, [], None
            acc_components = []

            # fall through to post-flush component rule

        # 4) Post-flush component capture
        if awaiting_components and last_msg_idx is not None:
            comp = _extract_component_from_line(line)
            if comp:
                messages[last_msg_idx]['component_lines'].append(comp)
            i += 1
            continue

        # default advance
        i += 1

    # 5) Trailing group at EOF
    if acc_bits:
        _flush_ixl_group(acc_label, acc_bits, acc_time,
                         messages, start_dt, end_dt,
                         acc_components)  # NEW

    # 6) Finalize: join component lines into a single string with newlines
    for msg in messages:
        lines_join = msg.get('component_lines', [])
        msg['component'] = '\r\n'.join(lines_join) if lines_join else ''
        msg.pop('component_lines', None)

    return messages


from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, PatternFill

def adjust_ixl_component_column(ws, header_row_idx: int = 2, start_row: int = 2,
                                ixl_component_col_index: int = 6,
                                min_width: int = 28, max_width: int = 60):
    """
    Enable wrap text and set a readable width for the IXL 'component' column.
    - Places each component on its own line has already been done in the data.
    - Here we wrap text and size the column based on the longest line (bounded).
    Parameters:
        ws: openpyxl worksheet
        header_row_idx: row where column headers are (the second row in your sheet)
        start_row: first data row (2 if header row 1 is merged header and row 2 is per-column headers)
        ixl_component_col_index: overall Excel column index for 'component' (default 6 = column F)
        min_width / max_width: bounds for final column width
    """
    col_letter = get_column_letter(ixl_component_col_index)

    # Turn on wrap + top-left alignment for all cells in that column
    for row in ws.iter_rows(min_row=start_row, max_row=ws.max_row,
                            min_col=ixl_component_col_index, max_col=ixl_component_col_index):
        for cell in row:
            cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

    # Estimate a good width from the longest single line across all cells in this column
    longest_line_len = 0
    for row in ws.iter_rows(min_row=start_row, max_row=ws.max_row,
                            min_col=ixl_component_col_index, max_col=ixl_component_col_index):
        for cell in row:
            val = str(cell.value) if cell.value is not None else ''
            # components are newline-separated; size by the longest line
            for ln in val.splitlines():
                if len(ln) > longest_line_len:
                    longest_line_len = len(ln)

    # Rough heuristic: 1 char ~ 1 unit; clamp between min_width and max_width
    target_width = max(min_width, min(max_width, longest_line_len + 2))
    ws.column_dimensions[col_letter].width = target_width


    # --- NEW: scale the row height to number of lines so all lines are visible ---
    # Heuristic: ~14 pts per text line + a little padding
    per_line_height = 14
    padding = 4
    
    fixed_height = 15  # Adjust as needed
    for r in range(start_row, ws.max_row + 1):
        ws.row_dimensions[r].height = fixed_height




def build_ixl_dataframe(ixl_file_path: str,
                        ws_entries: list[list],
                        start_time_str: str | None = None,
                        end_time_str: str | None = None,
                        max_diff_minutes: float = MAX_IXL_WS_DIFF_MINUTES) -> pd.DataFrame:
    """
    Build an IXL DataFrame row-aligned to ws_entries.
      * IXL time filter applied
      * exact payload + same type
      * |IXL - WS| <= max_diff_minutes
      * 'component' contains multi-line text (newline-separated) for Excel
    """
    row_count = len(ws_entries)
    aligned_empty = [{'time tag': '', 'msg type': '', 'data (hex)': '', 'direction': '', 'component': ''} for _ in range(row_count)]

    if not ixl_file_path or not os.path.isfile(ixl_file_path):
        return pd.DataFrame(aligned_empty, columns=["time tag", "msg type", "data (hex)", "direction", "component"])

    start_dt = end_dt = None
    try:
        if start_time_str:
            start_dt = parse_time_flexible(start_time_str)
        if end_time_str:
            end_dt = parse_time_flexible(end_time_str)
    except Exception:
        start_dt = end_dt = None

    ixl_msgs = parse_ixl_file(ixl_file_path, start_dt=start_dt, end_dt=end_dt)

    unmatched_ixl = [{
        'time_tag': m['time_tag'],
        'msg_type': m['msg_type'],
        'component': m.get('component', ''),              # already newline-joined
        'data_pretty': m['data_hex'],
        'data_norm': normalize_hex_no_spaces(m['data_hex']),
    } for m in ixl_msgs]

    rows = [{'time tag': '', 'msg type': '', 'data (hex)': '', 'direction': '', 'component': ''} for _ in range(row_count)]

    for i, ws_row in enumerate(ws_entries):
        if not ws_row or len(ws_row) < 4:
            continue

        ws_time_str = ws_row[0] or ''
        ws_msg_type = ws_row[1]
        ws_payload_pretty = ws_row[3]

        if not ws_msg_type or not ws_payload_pretty or not ws_time_str:
            continue
        if ws_msg_type not in ('Control (12 01)', 'Ind (12 8B)'):
            continue

        ws_payload_norm = normalize_hex_no_spaces(ws_payload_pretty)

        try:
            ws_dt = parse_time_flexible(ws_time_str)
        except Exception:
            continue

        found_idx = None
        for j, im in enumerate(unmatched_ixl):
            if im['msg_type'] != ws_msg_type:
                continue
            if im['data_norm'] != ws_payload_norm:
                continue
            try:
                ixl_dt = parse_time_flexible(im['time_tag'])
            except Exception:
                continue
            if minutes_abs(ws_dt, ixl_dt) <= float(max_diff_minutes):
                found_idx = j
                break

        if found_idx is not None:
            im = unmatched_ixl.pop(found_idx)
            rows[i]['time tag']   = im['time_tag']
            rows[i]['msg type']   = im['msg_type']
            rows[i]['data (hex)'] = im['data_pretty']

           # Creates a line with text "See more ..."
            component_text = im.get('component', '')
            if component_text:
                component_text = "See more ...\r\n" + component_text
            rows[i]['component'] = component_text

    return pd.DataFrame(rows, columns=["time tag", "msg type", "data (hex)", "direction", "component"])
# =================== END IXL parsing & matching (RULE-BASED COMPONENTS, NEWLINES) =====================

# IXL end
###################################################

def analyze_logs(ixl_file_path, log_file_path, pcap_file_path, ixl_excel_file_path, start_time_str, end_time_str, packetswitch_file_path, target_address, filename_suffix):

    # This is the filter applied
    target_address = target_address  # Target address to match in packets

    # Changes wanted for later
    # Add file to be rewritten
    # Allow user to name ending of ouput filename after log_packet_analysis
    output_file_path = f"log_packet_analysis_output_{filename_suffix or datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"    

    if os.path.exists(output_file_path) and is_file_open(output_file_path):
        messagebox.showerror("Error", "The output file is currently open. Please close it and try again.")
        return

    # Initialize storage for results and tracking
    cm_entries = []  # Communication Manager log entries
    ws_entries = []  # Wireshark log entries
    time_differences = []  # Time differences between logs
    last_two_sequences = []  # Last two hex sequences processed
    recent_sequence_times = {}  # Last timestamp for each hex sequence

    # As long as there is a communication manager log file given
    if log_file_path != False:
        
        # Parse start and end time inputs
        try:
            start_dt = parse_time_only(start_time_str)
            end_dt = parse_time_only(end_time_str)
        except Exception as e:
            messagebox.showerror("Error", f"Invalid time format: {e}")
            return

        # Read log and pcap files
        with open(log_file_path, 'r') as log_file:
            lines = log_file.readlines()
        packets = rdpcap(pcap_file_path)

        # Read packetswitch files using Beautiful soup
        html_text = ""
        if packetswitch_file_path.lower().endswith(".html"):
            with open(packetswitch_file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
                html_text = soup.get_text()
        
        elif packetswitch_file_path.lower().endswith(".docx"):
            # Before reading packetswitch file
            # Ensures packetswitch file is not open
            if is_file_open(packetswitch_file_path):
                messagebox.showerror("Error", "The Packetswitch file is currently open. Please close it and try again.")
                return

            try:
                html_text = "\n".join([para.text for para in Document(packetswitch_file_path).paragraphs])
            except:
                with open(packetswitch_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    html_text = f.read()


        # Process each line in the log file
        i = 0
        while i < len(lines):
            line = lines[i]
            time_match = time_pattern.search(line)
            if time_match:
                last_timestamp = time_match.group()
                try:
                    log_dt = parse_time_only(last_timestamp)
                except:
                    i += 1
                    continue
                # Only looks at entries between start and end time from user
                if not (start_dt <= log_dt <= end_dt):
                    i += 1
                    continue
            else:
                i += 1
                continue

            # Look for hex sequence pattern '02 02'
            if '02 02' in line:
                try:
                    index = line.index('02 02')
                    after = line[index + len('02 02'):].strip()
                    after_parts = after.split()
                    # This is in case the message number is '02'
                    if len(after_parts) >= 3 and after_parts[0] == '02':
                        msg_number = '02'
                        msg_type_raw = f"{after_parts[1]} {after_parts[2]}"
                    else:
                        # This is in all other cases a '02 02' message was found, but the message number is not '02'
                        pre_match = re.search(r'(\S{2})\s+02\s+02', line)
                        post_match = re.search(r'02\s+02\s+(\S{2})\s+(\S{2})', line)
                        msg_number = pre_match.group(1) if pre_match else None
                        msg_type_raw = f"{post_match.group(1)} {post_match.group(2)}" if post_match else None
                        
                        # Determine line type
                        line_type = "BASIC" if "BASIC" in line else "INFO" if "INFO" in line else ""

                    msg_type = msg_type_raw
                    if msg_type_raw == '12 8B':
                        msg_type = 'Ind (12 8B)'
                    elif msg_type_raw == '12 01':
                        msg_type = 'Control (12 01)'
                    elif msg_type_raw == '12 48':
                        msg_type = 'Recall (12 48)'
                    
                    hex_sequence = f"{msg_number} 02 02 {msg_type_raw}" if msg_number and msg_type_raw else None
                except:
                    i += 1
                    continue

                # Skip if hex sequence is in last two entries read
                if not hex_sequence or hex_sequence in last_two_sequences:
                    i += 1
                    continue

                # Skip if hex sequence was seen in last 5 seconds
                if hex_sequence in recent_sequence_times:
                    if (log_dt - recent_sequence_times[hex_sequence]) <= timedelta(seconds=30):
                        i += 1
                        continue
                # Convert hex sequence to bytes
                try:
                    search_bytes = bytes.fromhex(hex_sequence)
                except:
                    i += 1
                    continue

                # Find matching pcap timestamp
                pcap_ts_str, time_diff_str, _, ws_hex_data = find_pcap_time(search_bytes, log_dt, packets, target_address, pcap_file_path, msg_type_raw)

                # Does not print entries found with overflow
                if time_diff_str != "Found but overflow":
                    # Store results
                    hex_data = extract_hex_data(lines, i, msg_type_raw, line_type)
                    cm_entries.append([last_timestamp, msg_type, msg_number, hex_data, ''])
                    ws_entries.append([pcap_ts_str or '', '' if pcap_ts_str is None else msg_type, '' if pcap_ts_str is None else msg_number, ws_hex_data, ''])
                    time_differences.append([time_diff_str])

                # Update tracking
                last_two_sequences.append(hex_sequence)
                if len(last_two_sequences) > 2:
                    last_two_sequences.pop(0)
                recent_sequence_times[hex_sequence] = log_dt

            i += 1


    # SECOND ITERATION: Scan Wireshark for 02 02 codes not found in CM log (also looks for RF_ACK now)
    # Only iteration if no communication manager log file is given
    last_rf_ack_time = None
    if log_file_path == False:
        # Initialized empty list for found communication manager log times (will remain empty)
        found_cm_times = set()
    else:
        found_cm_times = set(entry[0] for entry in cm_entries)
    additional_ws_entries = []
    additional_cm_entries = []
    additional_time_differences = []

    if log_file_path == False:
        packets = rdpcap(pcap_file_path)
        # Filter packets based on start time and end time from user given there is no comm manager file
        try:
            start_dt = parse_time_only(start_time_str)
            end_dt = parse_time_only(end_time_str)
        except Exception as e:
            messagebox.showerror("Error", f"Invalid time format: {e}")
            return

        filtered_packets = []
        for pkt in packets:
            try:
                pkt_time = datetime.fromtimestamp(float(pkt.time))
                pkt_time_only = datetime.strptime(pkt_time.strftime('%H:%M:%S.%f')[:-3], '%H:%M:%S.%f')
                if start_dt <= pkt_time_only <= end_dt:
                    filtered_packets.append(pkt)
            except:
                continue
        packets = filtered_packets

        # Read packetswitch files
        html_text = ""
        if packetswitch_file_path.lower().endswith(".html"):
            with open(packetswitch_file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
                html_text = soup.get_text()
        
        elif packetswitch_file_path.lower().endswith(".docx"):  
            # Before reading packetswitch file
            if is_file_open(packetswitch_file_path):
                messagebox.showerror("Error", "The Packetswitch file is currently open. Please close it and try again.")
                return

            try:
                from docx import Document
                html_text = "\n".join([para.text for para in Document(packetswitch_file_path).paragraphs])
            except:
                with open(packetswitch_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    html_text = f.read()

        
    for pkt in packets:
        if Raw in pkt:
            data_bytes = pkt[Raw].load
            if is_valid_packet(pkt, data_bytes, target_address):
                hex_str = data_bytes.hex().upper()
                hex_parts = [hex_str[i:i+2] for i in range(0, len(hex_str), 2)]
                for i in range(len(hex_parts) - 3):

                    # Additional logic to detect '8C' messages (RF_ACK)
                    for j in range(len(hex_parts) - 4):
                        if hex_parts[j] == '8C':
                            fourth_pair = hex_parts[j + 4]
                            if fourth_pair in ['34', '38']:
                                rf_ack_time = datetime.fromtimestamp(float(pkt.time))
                                if last_rf_ack_time and (rf_ack_time - last_rf_ack_time) <= timedelta(seconds=30):
                                    continue
                                msg_type = 'RF_ACK'
                                # Clarify between inbound and outbound
                                if fourth_pair in ['34']:
                                    msg_type = 'RF_ACK (Inbound)'
                                elif fourth_pair in ['38']:
                                    msg_type = 'RF_ACK (Outbound)'
                                last_rf_ack_time = rf_ack_time
                                pcap_time = datetime.fromtimestamp(float(pkt.time)).strftime('%H:%M:%S.%f')[:-3]
                                additional_cm_entries.append(['Not found', '', '', '', ''])
                                additional_ws_entries.append([pcap_time, msg_type, '', '', ''])
                                additional_time_differences.append([''])
                                break
                        # Now look for '02 02'
                        if hex_parts[i] == '02' and hex_parts[i+1] == '02':

                            # Check if target address suffix is within 11 hex pairs before '02' (otherwise invalid 02 02)
                            suffix = target_address.hex()[-2:]
                            search_window = hex_parts[max(0, i-11):i]
                            if suffix not in search_window:
                                continue

                            # In case message number is 02  
                            if i > 0 and hex_parts[i+2] == '02':
                                msg_number = '02'
                                msg_type_raw = f"{hex_parts[i+3]} {hex_parts[i+4]}"
                                
                            else:
                                msg_number = hex_parts[i-1] if i > 0 else None
                                msg_type_raw = f"{hex_parts[i+2]} {hex_parts[i+3]}"

                            msg_type = msg_type_raw
                            
                            if msg_type_raw == '04 D0':
                                continue

                            if msg_type_raw == '12 8B':
                                msg_type = 'Ind (12 8B)'
                            elif msg_type_raw == '12 01':
                                msg_type = 'Control (12 01)'
                            elif msg_type_raw == '12 48':
                                msg_type = 'Recall (12 48)'
                            hex_sequence = f"{msg_number} 02 02 {msg_type_raw}" if msg_number and msg_type_raw else None


                            if log_file_path != False:
                                # Skip if this hex sequence was already processed in the first iteration
                                if hex_sequence in [f"{entry[2]} 02 02 {entry[1].split('(')[-1].strip(')')}" for entry in cm_entries if entry[0] != 'Not found']:
                                    continue

                            pcap_time = datetime.fromtimestamp(float(pkt.time)).strftime('%H:%M:%S.%f')[:-3]

                            # Skip if hex sequence is in last two entries
                            if not hex_sequence or hex_sequence in last_two_sequences:
                                continue

                            # Skip if hex sequence was seen in last 5 seconds
                            if hex_sequence in recent_sequence_times:
                                if (datetime.fromtimestamp(float(pkt.time)) - recent_sequence_times[hex_sequence]) <= timedelta(seconds=30):
                                    continue

                            # Update tracking
                            last_two_sequences.append(hex_sequence)
                            if len(last_two_sequences) > 2:
                                last_two_sequences.pop(0)
                            recent_sequence_times[hex_sequence] = datetime.fromtimestamp(float(pkt.time))

                            if pcap_time not in found_cm_times:
                                additional_cm_entries.append(['Not found', '', '', '', ''])
                                if log_file_path == False:
                                    
                                    fallback_hex = ' '.join([data_bytes.hex()[i:i+2] for i in range(0, len(data_bytes.hex()), 2)])
                                    ws_hex_data = extract_ws_hex_data(pcap_file_path, pcap_time, msg_type_raw, fallback_hex)

                                    additional_ws_entries.append([pcap_time, msg_type, msg_number, ws_hex_data, ''])

                                else:
                                    
                                    fallback_hex = ' '.join([data_bytes.hex()[i:i+2] for i in range(0, len(data_bytes.hex()), 2)])
                                    ws_hex_data = extract_ws_hex_data(pcap_file_path, pcap_time, msg_type_raw, fallback_hex)

                                    additional_ws_entries.append([pcap_time, msg_type, msg_number, ws_hex_data, ''])
                            
                                additional_time_differences.append([''])

    # Append new entries to original lists
    cm_entries.extend(additional_cm_entries)
    ws_entries.extend(additional_ws_entries)
    time_differences.extend(additional_time_differences)


    # Filter all entries based on user-defined start and end time
    filtered_entries = []
    for i, entry in enumerate(cm_entries):
        cm_time_str = entry[0]
        ws_time_str = ws_entries[i][0]

        try:
            # Parse timestamps
            cm_time = parse_time_only(cm_time_str) if cm_time_str != 'Not found' else None
            ws_time = parse_time_only(ws_time_str) if ws_time_str else None

            # Check if either timestamp is within range
            if (cm_time and start_dt <= cm_time <= end_dt) or (ws_time and start_dt <= ws_time <= end_dt):
                filtered_entries.append((entry, ws_entries[i], time_differences[i]))
        except:
            continue

    # Unpack filtered entries
    cm_entries = [e[0] for e in filtered_entries]
    ws_entries = [e[1] for e in filtered_entries]
    time_differences = [e[2] for e in filtered_entries]


    # Rebuild sorted entries based on CM time tag, placing 'Not found' entries before the next WS time
    # Step 1: Separate entries
    combined_entries = list(zip(cm_entries, ws_entries, time_differences))

    # Step 2: Sort entries with valid CM time tags
    valid_entries = [e for e in combined_entries if e[0][0] != 'Not found']
    valid_entries.sort(key=lambda x: parse_time_only(x[0][0]))

    # Step 3: Insert 'Not found' entries based on WS time tag
    did_not_find_entries = [e for e in combined_entries if e[0][0] == 'Not found']

    for entry in did_not_find_entries:
        ws_time_str = entry[1][0]
        try:
            ws_time = parse_time_only(ws_time_str)
        except:
            ws_time = datetime.max

        inserted = False
        for i, valid_entry in enumerate(valid_entries):
            try:
                next_ws_time = parse_time_only(valid_entry[1][0])
                if ws_time < next_ws_time:
                    valid_entries.insert(i, entry)
                    inserted = True
                    break
            except:
                continue
        if not inserted:
            valid_entries.append(entry)

    # Step 4: Unpack sorted entries
    cm_entries = [e[0] for e in valid_entries]
    ws_entries = [e[1] for e in valid_entries]
    time_differences = [e[2] for e in valid_entries]

    for entry in ws_entries:
        msg_type_raw = entry[1].split('(')[-1].strip(')') if '(' in entry[1] else entry[1]
        entry[3] = process_ws_hex_data(entry[3], msg_type_raw)



    # Create dataframes for Excel Output
    row_count = len(cm_entries)  # keep alignment across sections
    ixl_df = build_ixl_dataframe(ixl_file_path, ws_entries, start_time_str, end_time_str)

    # 2) The first Time Difference (between IXL and CM) should be empty
    diff2_df = pd.DataFrame({" ": [''] * row_count})

    # Should use this for ixl dataframe in the future
    # ixl_df = pd.DataFrame(ixl_entries, columns=["time tag", "msg type", "data (hex)", "direction", "component"])
    cm_df = pd.DataFrame(cm_entries, columns=["time tag", "msg type", "msg number", "data (hex)", "direction"])
    ws_df = pd.DataFrame(ws_entries, columns=["time tag", "msg type", "msg number", "data (hex)", "UDP/TCP"])
    diff_df = pd.DataFrame(time_differences, columns=[" "])


    # If packetswitch file is not provided or invalid, skip parsing
    if not packetswitch_file_path or not os.path.isfile(packetswitch_file_path):
        html_text = ''
    else:
        # Detect if packetswitch file contains 'Generic Report Results'
        is_generic_report = 'Generic Report Results' in html_text

        if not is_generic_report: # Not raw data packetswitch file
            
            # Packetswitch integration
            packetswitch_times, packetswitch_codes = [], []
            packetswitch_count = {}

            exit_bool = False
            for entry in ws_entries:
                msg_type_raw = entry[1].split('(')[-1].strip(')') if '(' in entry[1] else entry[1]
                if msg_type_raw not in ['12 8B', '12 01', '12 48']:
                    packetswitch_times.append('')
                    packetswitch_codes.append('')
                    continue
                pcap_time = entry[0]
                msg_type = entry[1]
                code = ''

                if pcap_time:
                    time_tag = pcap_time.split('.')[0]
                    key = (time_tag, msg_type)
                    count = packetswitch_count.get(key, 0)
                    idx = -1

                    for _ in range(count + 1):
                        idx = html_text.find(time_tag, idx + 1)
                        if idx == -1:
                            break

                    if idx != -1:
                        after = html_text[idx:]
                        underscore_idx = after.find('_')
                        if underscore_idx != -1 and len(after) > underscore_idx + 2:
                            raw_code = after[underscore_idx + 1:underscore_idx + 3]
                            if raw_code == "CR" and msg_type == "Control (12 01)":
                                code = "Control"
                            elif raw_code == "IR" and msg_type == "Ind (12 8B)":
                                code = "Ind"
                            elif raw_code == "R_" and msg_type == "Recall (12 48)":
                                code = "Recall"

                    if code:
                        packetswitch_times.append(time_tag)
                        packetswitch_codes.append(code)
                        packetswitch_count[key] = count + 1
                        exit_bool = True
                    else:
                        exit_bool = False
                # Determine if the tenth digit of the pcap timestamp is '9'
                # Some entrie where the pcap timestamp is '9' have a matching entry in the next second
                tenth_digit_is_nine = False
                
                try:
                    fractional = pcap_time.split('.')[-1]
                    if len(fractional) >= 1 and fractional[0] == '9':
                        tenth_digit_is_nine = True
                except:
                    pass


                # If no match found and tenth digit is 9, try again with +1 second
                if not code and tenth_digit_is_nine:
                    adjusted_time = (datetime.strptime(time_tag, '%H:%M:%S') + timedelta(seconds=1)).strftime('%H:%M:%S')
                    key = (adjusted_time, msg_type)
                    count = packetswitch_count.get(key, 0)
                    idx = -1
                    for _ in range(count + 1):
                        idx = html_text.find(adjusted_time, idx + 1)
                        if idx == -1:
                            break
                    if idx != -1:
                        after = html_text[idx:]
                        underscore_idx = after.find('_')
                        if underscore_idx != -1 and len(after) > underscore_idx + 2:
                            raw_code = after[underscore_idx + 1:underscore_idx + 3]
                            if raw_code == "CR" and msg_type == "Control (12 01)":
                                code = "Control"
                            elif raw_code == "IR" and msg_type == "Ind (12 8B)":
                                code = "Ind"
                            elif raw_code == "R_" and msg_type == "Recall (12 48)":
                                code = "Recall"
                    exit_bool = False
                    if code:
                        # packetswitch_times[-1] = adjusted_time
                        # packetswitch_codes[-1] = code
                        # packetswitch_count[key] = count + 1
                        # exit_bool = True
                        packetswitch_times.append(adjusted_time)
                        packetswitch_codes.append(code)
                        packetswitch_count[key] = count + 1
                        exit_bool = True
                        
                if exit_bool == False:
                    packetswitch_times.append('')
                    packetswitch_codes.append('')

        
        if is_generic_report: # Now if there is a raw data file
            packetswitch_times = []
            packetswitch_components = []
            packetswitch_codes = [] # Stores descriptive text
            packetswitch_data = []  # Matched hex data


            # Parse packetswitch lines differently for Generic Report Results
            ps_lines = html_text.strip().split('\n')
            parsed_ps_lines = []
            for line in ps_lines:
                # Capture time, descriptive text, and hex data
                match = re.match(r".*?(\d{2}:\d{2}:\d{2})\s+([A-Za-z()]+)\s.*?:\s*([0-9A-Fa-f ]+)$", line)
                if match:
                    time_tag = match.group(1)
                    description = match.group(2)  # Text after time tag and before next number
                    data = match.group(3).strip()
                    parsed_ps_lines.append((time_tag, description, data))

            # For each Wireshark entry, check for matching data
            for ws_entry in ws_entries:
                ws_time_str = ws_entry[0].split('.')[0]  # Extract HH:MM:SS
                try:
                    ws_time = datetime.strptime(ws_time_str, '%H:%M:%S')
                except:
                    ws_time = None

                # Normalize Wireshark data: remove spaces, uppercase, and trim last two hex digits
                ws_data = ws_entry[3].replace(" ", "").upper()

                found_match = False
                for time_tag, description, ps_data in parsed_ps_lines:
                    ps_data_clean = ps_data.replace(" ", "").upper()

                    # Check if packetswitch data starts with Wireshark trimmed data
                    if ps_data_clean.startswith(ws_data):
                        try:
                            ps_time = datetime.strptime(time_tag, '%H:%M:%S')
                            time_diff = abs((ps_time - ws_time).total_seconds()) if ws_time else None
                        except:
                            time_diff = None

                        # Ensure time difference is within ±1 second
                        if time_diff is not None and time_diff <= 1:
                            if ws_entry[1] == "RF_ACK (Outbound)" or ws_entry[1] == "RF_ACK (Inbound)":
                                continue
                            if description == 'Indic(RF)':
                                if ws_entry[1] != 'Ind (12 8B)':   
                                    continue
                                else:
                                    description = 'Ind'
                            
                            if description == 'Rf':
                                if ws_entry[1] != 'Control (12 01)':
                                    continue  # Skip this packetswitch line
                                else:
                                    description = 'Control'

                            if description == 'Indicate':
                                continue
                            
                            packetswitch_times.append(time_tag)
                            packetswitch_components.append('')
                            packetswitch_codes.append(description) # Store descriptive text
                            packetswitch_data.append(ps_data)
                            found_match = True
                            break

                # If no match found, append empty placeholders
                if not found_match:
                    packetswitch_times.append('')
                    packetswitch_components.append('')
                    packetswitch_codes.append('')
                    packetswitch_data.append('')


        if is_generic_report:
            # Create dataframe for Packetswitch (raw ps)
            packetswitch_df = pd.DataFrame({
                "time tag": packetswitch_times,
                "component": [''] * len(packetswitch_times),  # keep empty
                "data (hex)": packetswitch_data,        # hex data here
                "data type": packetswitch_codes
            })
        else:       
            # Create dataframe for Packetswitch (readable ps)
            packetswitch_df = pd.DataFrame({
                "time tag": packetswitch_times,
                "component": [''] * len(packetswitch_times),
                "data (hex)": [''] * len(packetswitch_times), # no hex data available
                "data type": packetswitch_codes
            })
    
    # If html_text is empty, create empty packetswitch dataframe
    if not html_text.strip():
        packetswitch_df = pd.DataFrame({
            "time tag": [''] * len(ws_entries),
            "component": [''] * len(ws_entries),
            "data (hex)": [''] * len(ws_entries),
            "data type": [''] * len(ws_entries)
        })


    # Combine all data into one dataframe with 'Number' column at the far left
    combined_df = pd.concat([
        pd.DataFrame({"Number": list(range(1, len(cm_df) + 1))}),  # New column added here
        ixl_df, diff2_df, cm_df, diff_df, ws_df, packetswitch_df
    ], axis=1)

    # Create header row for excel
    header_row = [
        " ",
        "IXL Log", "", "", "", "",
        "Time Difference",
        "Communication Manager Log", "", "", "", "",
        "Time Difference",
        "Wireshark Log", "", "", "", "",
        "Packetswitch Data", "", "", ""
    ]

    # Write to excel file
    with pd.ExcelWriter(output_file_path, engine="openpyxl") as writer:
        pd.DataFrame([header_row]).to_excel(writer, index=False, header=False)
        combined_df.to_excel(writer, index=False, startrow=1)

    # Format Excel file: merge headers and center align
    wb = load_workbook(output_file_path)
    ws = wb.active

    # Freeze top two rows
    ws.freeze_panes = "A3"
    
    # Merge headers
    ws.merge_cells(start_row=1, start_column=2, end_row=1, end_column=6)  # IXL Log
    ws.merge_cells(start_row=1, start_column=7, end_row=1, end_column=7)  # Time Diff
    ws.merge_cells(start_row=1, start_column=8, end_row=1, end_column=12)  # CM Log
    ws.merge_cells(start_row=1, start_column=13, end_row=1, end_column=13)  # Time Diff
    ws.merge_cells(start_row=1, start_column=14, end_row=1, end_column=18) # Wireshark
    ws.merge_cells(start_row=1, start_column=19, end_row=1, end_column=22) # Packetswitch

    # Center Align
    for col in [2, 7, 8, 13, 14, 19]:
        cell = ws.cell(row=1, column=col)
        cell.alignment = Alignment(horizontal='center', vertical='center')

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=22):
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center')

    
    # --- Enable wrap text for IXL component column (overall column F = 6) ---
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=6, max_col=6):
        for cell in row:
            # Align top-left so multiple lines read naturally; enable wrap
            cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

    # Auto-adjust column widths
    for col in ws.iter_cols(min_row=2):
        max_length = 0
        column = col[0].column
        column_letter = get_column_letter(column)
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column_letter].width = max_length + 2

    
    # After global auto-fit, override the IXL 'component' column (overall column F = 6):
    adjust_ixl_component_column(ws, header_row_idx=2, start_row=2, ixl_component_col_index=6,
                                min_width=28, max_width=60)

        
    # === Highlight message number jumps (Wireshark only) ===
    yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
    COL_WS_TYPE = 15
    COL_WS_NUM  = 16
    last_by_group = {}
    for r in range(2, ws.max_row + 1):
        ws_type_val = ws.cell(row=r, column=COL_WS_TYPE).value
        ws_num_val  = ws.cell(row=r, column=COL_WS_NUM).value
        if not ws_type_val or not ws_num_val:
            continue  # Skip empty WS rows
        raw_type = ws_type_val.split('(')[-1].strip(')') if '(' in ws_type_val else ws_type_val
        if raw_type.upper() == '08 42':
            continue  # Skip highlighting for 08 42
        group = _group_from_msg_type_cell(ws_type_val)
        try:
            current = int(str(ws_num_val).strip(), 16)
        except Exception:
            continue
        last = last_by_group.get(group)
        if last is None:
            last_by_group[group] = current
            continue
        expected = (last + 2) % 256
        if current != expected:
            ws.cell(row=r, column=COL_WS_NUM).fill = yellow_fill
        last_by_group[group] = current
    # === End highlight logic ===

    wb.save(output_file_path)

    # Notify user of success
    messagebox.showinfo("Success", "Output has been saved to: log_packet_analysis_output.xlsx\nThe Excel file is ready to be viewed.")

# UI functions for file selection
def browse_ixl_file():
    filename = filedialog.askopenfilename(title="Select IXL Log File", filetypes=[("Text Files", "*.txt")])
    ixl_file_entry.delete(0, tk.END)
    ixl_file_entry.insert(0, filename)

def browse_log_file():
    filename = filedialog.askopenfilename(title="Select Communication Manager Log File", filetypes=[("Text Files", "*.txt")])
    log_file_entry.delete(0, tk.END)
    log_file_entry.insert(0, filename)

def browse_pcap_file():
    filename = filedialog.askopenfilename(title="Select Wireshark PCAP File", filetypes=[("PCAP Files", "*.pcap")])
    pcap_file_entry.delete(0, tk.END)
    pcap_file_entry.insert(0, filename)

def browse_packetswitch_file():
    filename = filedialog.askopenfilename(
        title="Select Packetswitch File",
        filetypes=[("HTML or DOCX Files", "*.html *.docx")]
    )
    packetswitch_file_entry.delete(0, tk.END)
    packetswitch_file_entry.insert(0, filename)

def browse_ixl_excel_file():
    filename = filedialog.askopenfilename(title="Select IXL Excel File", filetypes=[("Excel Files", "*.xlsx")])
    ixl_excel_file_entry.delete(0, tk.END)
    ixl_excel_file_entry.insert(0, filename)

# First function ran after the user selects run_analysis
def run_analysis():
    has_log = True
    ixl_file = ixl_file_entry.get()
    log_file = log_file_entry.get()
    pcap_file = pcap_file_entry.get()
    packetswitch_file = packetswitch_file_entry.get()
    ixl_excel_file = ixl_excel_file_entry.get()
    start_time = start_time_entry.get()
    end_time = end_time_entry.get()
    target_address_hex = target_address_entry.get().replace('0', 'a').replace('.', '')

    # Ensure user does not enter an odd number of hexadecimal digits
    import string
    def is_valid_hex(s):
        return len(s) % 2 == 0 and all(c in string.hexdigits for c in s)

    if not is_valid_hex(target_address_hex):
        messagebox.showerror("Error", "Target address must be even-length and contain only hex characters.")
        return

    if not os.path.isfile(pcap_file):
        messagebox.showerror("Error", "Please select a valid PCAP file path.")
        return
    # Only wireshark log is necessary
    
    if not ixl_file:
        ixl_file = False
    if not log_file:
        log_file = False
    if not ixl_excel_file:
        ixl_excel_file = False
        
    analyze_logs(ixl_file, log_file, pcap_file, ixl_excel_file, start_time, end_time, packetswitch_file, bytes.fromhex(target_address_hex), filename_suffix_entry.get())

# ********** Everything above this point is functions and libraries **********

# UI layout
root = tk.Tk()
root.title("Log Packet Analysis Tool")

tk.Label(root, text="IXL Text File:").grid(row=0, column=0, sticky="w", padx=10, pady=5)
ixl_file_entry = tk.Entry(root, width=60)
ixl_file_entry.grid(row=0, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_ixl_file).grid(row=0, column=2, padx=10)

tk.Label(root, text="Communication Manager Log Text File:").grid(row=1, column=0, sticky="w", padx=10, pady=5)
log_file_entry = tk.Entry(root, width=60)
log_file_entry.grid(row=1, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_log_file).grid(row=1, column=2, padx=10)

tk.Label(root, text="Wireshark PCAP File:").grid(row=2, column=0, sticky="w", padx=10, pady=5)
pcap_file_entry = tk.Entry(root, width=60)
pcap_file_entry.grid(row=2, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_pcap_file).grid(row=2, column=2, padx=10)

tk.Label(root, text="Packetswitch Data File:").grid(row=3, column=0, sticky="w", padx=10, pady=5)
packetswitch_file_entry = tk.Entry(root, width=60)
packetswitch_file_entry.grid(row=3, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_packetswitch_file).grid(row=3, column=2, padx=10)

tk.Label(root, text="IXL Excel File:").grid(row=4, column=0, sticky="w", padx=10, pady=5)
ixl_excel_file_entry = tk.Entry(root, width=60)
ixl_excel_file_entry.grid(row=4, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_ixl_excel_file).grid(row=4, column=2, padx=10)

tk.Label(root, text="Start Time (HH:MM:SS.sss):").grid(row=5, column=0, sticky="w", padx=10, pady=5)
start_time_entry = tk.Entry(root, width=20)
start_time_entry.grid(row=5, column=1, sticky="w", padx=10)

tk.Label(root, text="End Time (HH:MM:SS.sss):").grid(row=6, column=0, sticky="w", padx=10, pady=5)
end_time_entry = tk.Entry(root, width=20)
end_time_entry.grid(row=6, column=1, sticky="w", padx=10)

tk.Label(root, text="Target Address:").grid(row=7, column=0, sticky="w", padx=10, pady=5)
target_address_entry = tk.Entry(root, width=20)
target_address_entry.grid(row=7, column=1, sticky="w", padx=10)

tk.Label(root, text="Output suffix after 'log_packet_analysis_output':").grid(row=8, column=0, sticky="w", padx=10, pady=5)
filename_suffix_entry = tk.Entry(root, width=20)
filename_suffix_entry.grid(row=8, column=1, sticky="w", padx=10)

tk.Button(root, text="Run Analysis", command=run_analysis, bg="green", fg="white").grid(row=9, column=1, pady=20)

tk.Label(root, text="You will receive a message that the Excel file is ready to be viewed. This may take a few minutes", fg="blue").grid(row=10, column=0, columnspan=3, pady=10)

# Start the UI event loop
root.mainloop()
