"""
Log Packet Analysis Tool
Author: Jake Chinchar
Last Edited: 12/4/2025

Purpose
-------
Analyze Communication Manager logs, Wireshark PCAP files, and Packetswitch HTML/DOCX reports
and correlate message flows. Produces a formatted Excel report summarizing findings.

Wireshark Filter Example: 
((atcsl3.srce_addr contains 71:a3:a7:8a:36) || (atcsl3.dest_addr contains 71:a3:a7:8a:36)) && !(_ws.col.protocol == "ATCS/TCP")

For Git
cd "folder path"
git status
git add LogAnalyzerNew.py
git commit -m "Describe changes"
git push origin main

External packages:
    pip install scapy pandas openpyxl beautifulsoup4 python-docx
Built-in modules:
    tkinter, datetime, re, os, subprocess, json

Wireshark Filter Example
------------------------
((atcsl3.srce_addr contains 71:a3:a7:8a:36) (atcsl3.dest_addr contains 71:a3:a7:8a:36))
&& !(_ws.col.protocol == "ATCS/TCP")

Future Improvements 
-----------------------------------
1. Add date filtering
2. Discover addresses in Wireshark automatically for selection
3. Notify if incorrect input file type
4. Progress bar for long runs
5. Validate ATCS address from CM log against user input (when CM file provided)
6. Eliminate need for a Wireshark file
7. Prevent invalid output filenames (e.g., '/', '?')


High-Level Architecture
-----------------------
The script is a single-file GUI application (Tkinter) composed of five layers:

1) GUI & Input Layer
   - Tkinter window with file selectors and text fields.
   - Collects paths and parameters; performs basic validation.

2) Parsing Layer
   - CM log parser: scans for '02 02' message sequences and extracts message number
     and type, plus a payload trimmed by a custom length-byte rule.
   - Wireshark reader: loads PCAP via Scapy; uses heuristics and TShark JSON
     to find the exact hex payload for a specific timestamp.
   - IXL parser: groups adjacent CTL/IND lines (±0.1s), reverses bits per row, builds
     a hex payload, and attaches "component" lines (multi-line text).
   - Packetswitch parser: supports two formats:
       • Readable reports ("Control", "Ind", "Recall" via time-tag + code)
       • "Generic Report Results" with raw hex payloads

3) Correlation Layer
   - CM <-> Wireshark correlation by message number/type and +-8 minutes around CM time tag.
   - Wireshark-only iteration when CM log is not provided: finds additional 02 02 messages
     and RF_ACK frames (All RF_ACK messages have this pattern in the hexstream (8C ... 34/38). 
   - IXL <-> Wireshark alignment: same message type, exact payload match, and a configurable
     max time difference (default 8 minutes).
   - Packetswitch <-> Wireshark association using HH:MM:SS time tags (with a special case
     for fractional '9' → +1 second retry) or raw-data prefix match for Generic reports.
     Can remove the the fractional 9 because the +1 second retry handles it

4) Data Modeling Layer
   - Builds aligned Pandas DataFrames for: IXL, CM, WS, Packetswitch, and two Time
     Difference columns.
   - Rows are kept in sync to allow easy cross-reference in Excel.

5) Output & Formatting Layer
   - Writes Excel via openpyxl:
     • Merged top headers for each section
     • Freeze panes (A3)
     • Center alignment in most columns; top-left wrap for IXL “component”
     • Auto-fit with special width rules for the component column
     • Highlights WS message-number anomalies by group
        - Some messages like 12 8B and 78 8B share message numbers

Inputs (GUI Fields)
-------------------
1. IXL Text File (optional)
   - Path to a text log containing CTL/IND entries with time tags and bit strings.
2. Communication Manager Log Text File (optional)
   - Path to a text log where '02 02' sequences appear around message number/type.
3. Wireshark PCAP File (required)
   - Path to a .pcap file (Scapy-readable).
4. Packetswitch Data File (optional, but necessary to read packetswitch and IXL)
   - Path to a .html or .docx file containing either readable report text or raw packet lines.
5. IXL Excel File (optional, currently not used)
   - Path to .xlsx (placeholder retained; no functional read performed).
6. Start Time (HH:MM:SS.sss)
   - Inclusive lower bound (flexible parser also accepts .ss or .sss).
7. End Time (HH:MM:SS.sss)
   - Inclusive upper bound.
8. Target Address (hex)
   - Hex string representing device/address of interest (even length; no dots).
   - NOTE: The script replaces '0' with 'a' and removes '.' before conversion.
9. Output Suffix
   - Appended to the generated Excel file name: "log_packet_analysis_output_<suffix>.xlsx".
   - If empty, timestamp is used.

Expected File Formats
---------------------
• CM Log (text)
  - Contains timestamp lines (HH:MM:SS.ff) and hex flows.
  - The parser looks for '02 02' sequences and infers:
      - Message number: either the token before '02 02' or literal '02' special case.
      - Message type raw: two hex pairs after '02 02'.
  - Payload extraction rule:
      1) After '02 02 12 8B' or '02 02 12 01', collect hex tokens across subsequent lines
         until a TX/RX boundary is reached.
      2) Remove first two hex pairs.
      3) Use the next hex pair as a length byte:
         • Convert to 8-bit binary, drop the first 2 bits → remaining bits are payload length
           in hex pairs.
      4) Remove the length byte itself and trim the payload to that length.

• Wireshark PCAP
  - Filtered by Scapy: packet must contain target address bytes and not be TCP.
  - Correlation window around the CM timestamp is ±8 minutes (CM tags can drift).
  - When locating exact payload for a WS timestamp:
      - TShark (if available) is invoked to produce JSON; we search recursively for
        'indication_bits' (12 8B) or 'control_bits' (12 01) showname fields.
      - Fallback to raw Scapy bytes if TShark parse fails.

• IXL Text
  - Lines formatted as (examples):
      "11-20 14:32:07.45 ... CTL(12-34): 10101100"
      "11-20 14:32:07.46 ... IND(56-78): 01010101"
  - Grouping rule:
      • Merge up to next two CTL/IND rows within ±0.1 seconds and same label.
      • While a block is “open”, attach component lines whose time tags are within ±0.1 seconds.
      • After a block is flushed, component lines can continue attaching to the last Control/Ind,
        until the next block begins.
  - Bit-to-hex:
      • Per-row bit strings are reversed (bit order) and converted to hex bytes, then
        joined across the group.
  - Output includes a “component” field (newline-separated text).

• Packetswitch (HTML/DOCX)
  - Readable report: we find HH:MM:SS tags and look for code near an underscore:
      • "CR" → "Control", "IR" → "Ind", "R_" → "Recall"
    Special case:
      • If the tenth fractional digit of the WS time is '9', retry with +1 second.
  - Generic Report Results:
      • Lines with time, description, and hex payload.
      • We match Packetswitch hex that starts with the (processed) WS hex payload,
        within ±1 second around HH:MM:SS.

Processing & Flow (Step-by-Step)
--------------------------------
A) GUI → Input Collection
   1) User selects files and enters parameters.
   2) Basic validation:
      - PCAP path must exist.
      - Target address must be even-length valid hex.
      - Output file cannot be open/locked.

B) If CM log provided (primary correlation pass)
   3) Read CM lines and PCAP packets.
   4) Limit CM entries to Start/End time range.
   5) For each CM line with '02 02':
      - Extract message number and raw type (12 8B/12 01/12 48 etc.).
      - Build hex sequence string: "<msg_num> 02 02 <type>" → `search_bytes`.
      - Find matching WS packet timestamp within ±8 minutes via Scapy/TShark.
      - Extract CM payload via length-byte rule.
      - Record CM row, the WS row (timestamp/type/number/payload), and time difference.

   Rules
      - Skip if the same hex sequence appears in the last two processed entries.
      - Skip if the same hex sequence appeared in the last 30s (per CM time).

C) Second pass: Wireshark-only discovery (if CM log not provided)
   6) Filter PCAP packets by user Start/End time.
   7) For each valid WS packet payload:
      - Detect RF_ACK: '8C' … 5th nibble '34' (Inbound) or '38' (Outbound) within 5s.
      - Detect '02 02' sequences near target address suffix; derive msg number/type.
      - Build WS entries; CM side is "Not found".

D) Alignment & Sorting
   8) Filter combined entries by user time range (either CM or WS time inside window).
   9) Sort entries:
      - Valid CM time-tagged rows sorted by CM time.
      - "Not found" rows placed before the next WS time tag in the sorted list.

E) Post-processing / Normalization
   10) Trim WS payload via `process_ws_hex_data()`:
       - After '02 02 12 8B' or '02 02 12 01', drop 4 pairs and, for '12 8B', drop the last pair.
   11) Build aligned DataFrames:
       - IXL DataFrame row-aligned to WS (same type + exact payload, ≤ max diff minutes).
       - Place “See more …” line before multi-line IXL component text.
   12) Build Packetswitch DataFrame (readable or generic):
       - Provide “data type” (Control/Ind/Recall) or matched raw hex, plus time/component.

F) Excel Generation
   13) Concatenate sections into one sheet with a "Number" column.
   14) Write a merged header row and data rows starting at row 2.
   15) Freeze panes at A3; center-align most cells; wrap text and widen IXL "component" column.
   16) Highlight WS message-number anomalies by group (expected hex increment +2 modulo 256).

Outputs
-------
• Excel workbook: "log_packet_analysis_output_<suffix or timestamp>.xlsx"
  - Sheet columns (left → right):
      Number |
      IXL: time tag, msg type, data (hex), direction, component |
      Time Difference (empty column placeholder) |
      Communication Manager: time tag, msg type, msg number, data (hex), direction |
      Time Difference (CM↔WS minutes/label) |
      Wireshark: time tag, msg type, msg number, data (hex), UDP/TCP |
      Packetswitch: time tag, component (empty or text), data (hex) (generic only), data type

Status & Messaging
------------------
• On success: a GUI message indicates the Excel file is ready.
• On failure: GUI error messages for invalid times, locked files, or missing PCAP.

Key Assumptions & Rules
-----------------------
1) CM time tags are approximate; WS correlation allows ±8 minutes drift.
2) Target address must appear in packet payload; TCP packets are excluded by today’s rule.
3) CM payload length is derived from a specialized length-byte bit rule (drop first 2 bits).
4) IXL groups CTL/IND lines within ±0.1 seconds and reverses bits per row before hex conversion.
5) Prevent duplicate reporting:
   - Skip identical sequences seen in the last two entries.
   - Skip sequences within 30 seconds of last occurrence.
6) Packetswitch matching is heuristic:
   - Readable: time-tag and suffix code detection; special case for fractional '9'.
   - Generic: raw hex starts-with comparison within ±1 second.
7) Excel presentation is fixed and designed for cross-reading:
   - Merged headers, frozen panes, component text wrapping, and anomaly highlighting.


------------
External
  • scapy               – PCAP reading
  • pandas              – DataFrames/Excel writing (with openpyxl)
  • openpyxl            – Excel formatting & workbook edits
  • beautifulsoup4      – HTML parsing
  • python-docx         – DOCX text extraction
  • TShark (CLI)        – JSON output for precise Wireshark payload fields
Built-in
  • tkinter, datetime, re, os, subprocess, json

Typical Run (Example)
---------------------
1) Select PCAP, optionally CM and IXL logs, optionally Packetswitch (HTML/DOCX).
2) Enter Start/End times (e.g., 4:00:00 → 5:00:00) and Target Address hex.
3) Provide output suffix (e.g., "A_Line_12").
4) Click "Run Analysis".
5) Wait for the success message; open the generated Excel file.

Troubleshooting Tips
--------------------
• If the Excel output is not created:
  - Ensure the output file is closed (some editors lock the file).
• If The Excel Sheet is empty:
  - Confirm Start/End times are in a valid format and cover actual data windows (1pm = 13:00:00)
• Ensure there is an output suffix and the target address is always correct
"""

# ========================
# Imports
# ========================
import os                                               # Operating system functions
import re                                               # Pattern matching & parsing
import json                                             # Used to parse JSON output from tshark and handle structured data
import subprocess                                       # Used to run external commands like 'tshark' for extracting Wireshark data
from datetime import datetime, timedelta                # For dates and times

import tkinter as tk                                    # For GUI
from tkinter import filedialog, messagebox              # filedialog for opening/saving files, messagebox for error alerts and success notifs

import pandas as pd                                     # For dataframes
from bs4 import BeautifulSoup                           # To read HTML files
from docx import Document                               # To read .docx
    
from scapy.all import rdpcap, Raw                       # Read packets from pcap file
from scapy.layers.inet import TCP                       # To identify TCP packets

from openpyxl import load_workbook                      # To write to and modify excel files   
from openpyxl.styles import Alignment, PatternFill      # Used to set cell alignment in Excel (center, left, wrap text) 
from openpyxl.utils import get_column_letter            # Converts numeric column index to Excel letters (e.g., 1 -> 'A') for formatting


# ========================
# Constants & Regex
# ========================

# Expression to extract time tags from log lines (e.g., 12:34:56.78)
TIME_PATTERN = re.compile(r"\b\d{2}:\d{2}:\d{2}\.\d{2}\b")

IXL_LABEL_TO_WS_TYPE = {
    "CTL": "Control (12 01)",
    "IND": "Ind (12 8B)",
}
IXL_LINE_RE = re.compile(
    r"^(?:\w{3}\s+)?"  # Optional weekday (e.g., Wed)
    r"(?P<md>\d{2}-\d{2}(?:-\d{4})?)\s+"  # MM-DD or MM-DD-YYYY
    r"(?P<hms>\d{2}:\d{2}:\d{2}\.\d{2}).*?\b"
    r"(?P<label>CTL|IND)\(\d{1,3}-\d{1,3}\):\s+(?P<bits>[01]{8})"
)
IXL_END_EXECUTE_RE = re.compile(r"Execute issued", re.IGNORECASE)
MAX_IXL_WS_DIFF_MINUTES = 8


# ========================
# Small utilities
# ========================

def parse_time_only(ts_str: str) -> datetime:
    """
    Parse 'HH:MM:SS.f' or 'HH:MM:SS.ff' (add .0 if missing fractional).
    Converts a time string to a datetime object
    Hours:Minutes:Seconds.Decimal
    """
    if "." not in ts_str:
        ts_str += ".0"
    return datetime.strptime(ts_str, "%H:%M:%S.%f")


def parse_time_flexible(ts_str: str) -> datetime:
    """Parse time allowing 2–3 fractional digits; normalize to milliseconds."""
    if not ts_str:
        raise ValueError("Empty time string")
    if "." in ts_str:
        hms, frac = ts_str.split(".", 1)
        if len(frac) == 2:
            ts_norm = f"{hms}.{frac}0"
        elif len(frac) >= 3:
            ts_norm = f"{hms}.{frac[:3]}"
        else:
            ts_norm = f"{hms}.{frac.ljust(3, '0')}"
    else:
        ts_norm = ts_str + ".000"
    return datetime.strptime(ts_norm, "%H:%M:%S.%f")


def format_timedelta(td: timedelta) -> str:
    """
    Format timedelta into HH:MM:SS.ms (2-digit ms).
    Formats a timedelta object into a readable string (HH:MM:SS.ms)
    """
    total_seconds = int(td.total_seconds())
    milliseconds = int(td.microseconds / 10000)
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours:02}:{minutes:02}:{seconds:02}.{milliseconds:02}"


def minutes_abs(dt_a: datetime, dt_b: datetime) -> float:
    """Absolute minutes between times ignoring date component"""
    a = dt_a.time(); b = dt_b.time()
    to_secs = lambda t: t.hour * 3600 + t.minute * 60 + t.second + t.microsecond / 1_000_000
    return abs(to_secs(a) - to_secs(b)) / 60.0


def is_file_open(filepath: str) -> bool:
    """
    Return True if the file appears to be open/locked (cannot rename to itself)
    This script does not write over output files that are open. The user can close the file or enter a different name for the output file in the UI
    """
    try:
        os.rename(filepath, filepath)
        return False
    except Exception:
        return True

# ========================
# Wireshark helpers
# ========================

def is_valid_packet(pkt, data_bytes: bytes, target_address: bytes) -> bool:
    """Packet contains target_address and is not TCP (may not want just TCP in the future)"""
    return target_address in data_bytes and not TCP in pkt


def extract_ws_hex_data(pcap_path: str, ws_time_tag: str, msg_type_raw: str, fallback_hex: str = "") -> str:
    """
    Extract a specific hex payload from a Wireshark PCAP using tshark JSON output, matching
    exactly on a Wireshark-formatted timestamp ('HH:MM:SS.mmm')

    msg_type_raw expected values:
        '12 8B' -> look for 'indication_bits'
        '12 01' -> look for 'control_bits'

    Returns space-separated uppercase hex pairs, or fallback_hex if not found
    """

    """
    Map the raw message type to the corresponding label name that appears in
    tshark's JSON "showname" fields. If type isn't recognized, exits early
    """
    if msg_type_raw == "12 8B":
        wanted_label = "indication_bits"
    elif msg_type_raw == "12 01":
        wanted_label = "control_bits"
    else:
        return ""  # Unknown type

    try:
        """Phase 1: Narrow the search window to the matching second"""
        second_str = ws_time_tag.split(".")[0]  # HH:MM:SS
        display_filter = f'frame.time contains "{second_str}"'

        """
        Run tshark to produce JSON for frames in that second
        -r <pcap> : read file
        -Y <filter> : display filter (not capture filter)
        -T json : output as JSON
        """
        proc = subprocess.run(
            ["tshark", "-r", pcap_path, "-Y", display_filter, "-T", "json"],
            capture_output=True,
            text=True,
            check=True,
        )
        frames = json.loads(proc.stdout)

        """
        Phase 2: Find the single frame that exactly matches 'ws_time_tag'
        Recurse through dicts/lists
        Read each frame, get epoch time → convert to HH:MM:SS.mmm, compare
        """
        for fr in frames:
            layers = fr.get("_source", {}).get("layers", {})
            t_epoch_list = layers.get("frame", {}).get("frame.time_epoch", [])
            if not t_epoch_list:
                continue
            try:
                """Convert epoch to float → datetime → format 'HH:MM:SS.mmm'"""
                t_epoch = float(t_epoch_list[0])
                candidate = datetime.fromtimestamp(t_epoch).strftime("%H:%M:%S.%f")[:-3]
            except Exception:
                continue
            if candidate != ws_time_tag:
                continue

            """
            Phase 3: Read JSON tree to locate the value tied to `wanted_label` 
            Look specifically for a field with 'showname' that starts with '<wanted_label>: '.
            """
            
            def find_val(node):
                if isinstance(node, dict):
                    if "showname" in node and isinstance(node["showname"], str):
                        prefix = f"{wanted_label}: "
                        if node["showname"].startswith(prefix):
                            return node["showname"][len(prefix) :].strip()
                    for v in node.values():
                        val = find_val(v)
                        if val:
                            return val
                elif isinstance(node, list):
                    for item in node:
                        val = find_val(item)
                        if val:
                            return val
                return None

            val = find_val(layers)
            if val:
                """
                Phase 4: Normalize to hex pairs 
                Remove any non-hex characters, then group into uppercased byte pairs
                """
                raw = re.sub(r"[^0-9A-Fa-f]", "", val)
                return " ".join(raw[i : i + 2].upper() for i in range(0, len(raw), 2))
    except Exception:
        pass
    return fallback_hex
    """ Didn't find a precise match """


def find_pcap_time(search_bytes: bytes, log_dt: datetime, packets, target_address: bytes, pcap_file_path: str, msg_type_raw: str):
    """Find first matching packet with same msg type & number within ±8 minutes of CM time tag"""
    for pkt in packets:
        if Raw in pkt and target_address in pkt[Raw].load:
            data_bytes = pkt[Raw].load
            if search_bytes in data_bytes and target_address in data_bytes and not TCP in pkt:
                pcap_dt = datetime.fromtimestamp(float(pkt.time))
                pcap_ts_str = pcap_dt.strftime("%H:%M:%S.%f")[:-3]
                try:
                    pcap_time_only = datetime.strptime(pcap_ts_str, "%H:%M:%S.%f")
                    time_diff = pcap_time_only - log_dt
                    minutes_diff = time_diff.total_seconds() / 60
                    if pcap_time_only < (log_dt - timedelta(minutes=8)):
                        continue
                    if -8 <= minutes_diff <= 8: # May need to be changed, but for now an 8 minute difference or less is valid
                        fallback_hex = " ".join(
                            [data_bytes.hex()[i : i + 2] for i in range(0, len(data_bytes.hex()), 2)]
                        )
                        ws_hex_data = extract_ws_hex_data(pcap_file_path, pcap_ts_str, msg_type_raw, fallback_hex)
                        return pcap_ts_str, format_timedelta(time_diff), True, ws_hex_data
                    else:
                        """
                        The wireshark message found does not correspond to the communication manager message with the same msg number and type
                        Messages with "Found but overflow" are not going to be printed to the excel file
                        There is a finite number of message numbers, so this states that there was wrapping of the message number
                        """
                        return pcap_ts_str, "Found but overflow", False, ""
                except Exception:
                    continue
    """If no messages or overflow were found, then the message is "LOST", which means that the communication manager log has this message, but the wireshark did not"""
    return None, "LOST", False, ""


# ========================
# Hex processing
# ========================

def _normalize_group_from_raw(msg_type_raw: str) -> str:
    s = (msg_type_raw or "").upper().strip()
    if s in ("12 8B", "78 8B"):
        return "IND"
    if s in ("12 01", "12 48"):
        return "CTL_RECALL"
    if s in ("08 82", "08 83"):
        return "GROUP_0882_0883"
    return s


def _group_from_msg_type_cell(cell_val: str) -> str:
    val = (cell_val or "").strip()
    if "(" in val and ")" in val:
        raw = val.split("(")[-1].strip(")")
    else:
        raw = val
    return _normalize_group_from_raw(raw)


def process_ws_hex_data(ws_hex_data: str, msg_type_raw: str) -> str:
    """Trim Wireshark hex payload per known pattern"""
    if not ws_hex_data:
        return ws_hex_data
    parts = ws_hex_data.strip().split()
    idx = None
    for i in range(len(parts) - 3):
        if parts[i : i + 4] == ["02", "02", "12", "8b"] or parts[i : i + 4] == ["02", "02", "12", "01"]:
            idx = i + 4
            break
    if idx is None:
        return ws_hex_data
    trimmed = parts[idx:]
    trimmed = trimmed[4:] if len(trimmed) > 4 else []

    """ If the message type is 12 8b, then the last hexadecimal pair needs to be removed """
    if msg_type_raw == "12 8B" and len(trimmed) > 0:
        trimmed = trimmed[:-1]
    return " ".join(trimmed)


def extract_hex_data(lines, start_index: int, msg_type_raw: str, line_type: str) -> str:
    """
    After finding '02 02 12 8B' or '02 02 12 01', remove first 2 hex pairs
    Take next (third) hex pair, convert to binary; remove first 2 bits; remaining bits -> length
    Trim collected hex pairs to match this length
    """
    collected = []
    found_sequence = False
    sequence_patterns = [["02", "02", "12", "8B"], ["02", "02", "12", "01"]]
    i = start_index
    while i < len(lines):
        tokens = lines[i].strip().split()
        joined = " ".join(tokens).upper()
        if " TX" in joined or " RX" in joined or "TX:" in joined or "RX:" in joined:
            break
        if not found_sequence:
            for j in range(len(tokens) - 3):
                if tokens[j : j + 4] in sequence_patterns:
                    found_sequence = True
                    collected.extend([t for t in tokens[j + 4 :] if re.fullmatch(r"[0-9A-Fa-f]{2}", t)])
                    break
        else:
            collected.extend([t for t in tokens if re.fullmatch(r"[0-9A-Fa-f]{2}", t)])
        i += 1
    if not found_sequence or len(collected) < 3:
        return ""
    collected = collected[2:]
    length_hex = collected[0]
    length_bin = bin(int(length_hex, 16))[2:].zfill(8)
    length_bin_trimmed = length_bin[2:]
    length_val = int(length_bin_trimmed, 2)
    collected = collected[2:]
    if len(collected) > length_val:
        collected = collected[:length_val]
    return " ".join(collected)


# ========================
# IXL parsing & matching (rule-based components)
# ========================

"""
Future: move this section to its own file for better clarity

IXL parsing & matching (RULE-BASED COMPONENTS, NEWLINES)
 - Groups CTL/IND rows (per-row bit reversal).
 - Applies IXL start/end time filter.
 - Matches to Wireshark by same msg type + exact payload + |delta(t)| <= 12 minutes.
 - COMPONENT RULE:
     For any line containing " -- ", take the substring AFTER " -- ".
     If it does NOT start with IND/CTL/Execute (case-insensitive), capture it as a component line
     and attach it to the most recently completed Control or Ind message.
     Place each captured component on its own line within the same Excel cell.
"""

def reverse_bits_to_hex(bits_8: str) -> str:
    """ IXL hex data is reversed in the file"""
    rev = bits_8[::-1]
    return f"{int(rev, 2):02X}"


def bits_list_to_hex(bits_list) -> str:
    return " ".join(reverse_bits_to_hex(b) for b in bits_list)


def normalize_hex_no_spaces(s: str) -> str:
    return (s or "").replace(" ", "").upper()


def _extract_component_from_line(line: str):
    """Extract component text (non IND/CTL/Execute) after '--' or after timestamp."""
    anchor = line.find(" -- ")
    if anchor != -1:
        after = line[anchor + len(" -- ") :].strip()
    else:
        m = re.match(
            r"^(?:\w{3}\s+)?"               # optional weekday, e.g., 'Wed '
            r"\d{2}-\d{2}(?:-\d{4})?\s+"    # MM-DD or MM-DD-YYYY
            r"\d{2}:\d{2}:\d{2}\.\d{2}\s+"  # HH:MM:SS.ff
            r"(?P<rest>.+)$",               # the remainder is candidate component text
            line,
        )
        if not m:
            return None
        after = m.group("rest").strip()
    if not after:
        return None
    """ Ignore non-component starters """
    head = after[:16].lower()
    if head.startswith("ind(") or head.startswith("ctl(") or head.startswith("execute"):
        return None
    return after


def _flush_ixl_group(acc_label, acc_bits, acc_time, out_msgs, start_dt=None, end_dt=None, acc_components=None):
    """Append flushed message and return its index, or None if filtered out by time/type."""
    if not acc_label or not acc_bits:
        return None
    ws_type = IXL_LABEL_TO_WS_TYPE.get(acc_label.upper())
    if not ws_type:
        return None
    include = True
    if start_dt or end_dt:
        try:
            t = parse_time_flexible(acc_time)
            if start_dt and t.time() < start_dt.time():
                include = False
            if end_dt and t.time() > end_dt.time():
                include = False
        except Exception:
            include = False
    if not include:
        return None
    data_hex = bits_list_to_hex(acc_bits)
    comp_lines = acc_components or []
    out_msgs.append(
        {
            "time_tag": acc_time or "",
            "msg_type": ws_type,
            "data_hex": data_hex,
            "_data_norm": normalize_hex_no_spaces(data_hex),
            "component_lines": comp_lines,
        }
    )
    return len(out_msgs) - 1


def parse_ixl_file(ixl_file_path: str, start_dt=None, end_dt=None):
    """
    Parse IXL into grouped CTL/IND messages (per-row bit reversal).

    Merge rule:
      - If there is a raw data line (IND/CTL) in the next 2 lines whose
        time tag is within +-0.1 seconds of the first raw data line's time tag
        (and label matches), treat it as part of the same raw data block.

    Component rule:
      - While accumulating a block, if a non-data line's time tag is within
        +-0.1 seconds of the block's start time, attach it as a component
        to the same block (do NOT flush).
      - After a block is flushed, subsequent component lines are still
        attached by the existing "awaiting_components" behavior until
        the next CTL/IND row starts.
    """
    if not ixl_file_path or not os.path.isfile(ixl_file_path):
        return []

    def _within_point_one_seconds(t1_str: str, t2_str: str) -> bool:
        try:
            t1 = parse_time_flexible(t1_str)
            t2 = parse_time_flexible(t2_str)
            return abs((t2 - t1).total_seconds()) <= 0.1
        except Exception:
            return False

    def _time_from_any_line(s: str):
        m = TIME_PATTERN.search(s)
        return m.group(0) if m else None

    messages = []
    acc_label = None
    acc_bits = []
    acc_time = None
    acc_components = []
    awaiting_components = False
    last_msg_idx = None

    with open(ixl_file_path, "r", encoding="utf-8", errors="ignore") as f:
        lines = f.readlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip("\n")
        m = IXL_LINE_RE.search(line)
        if m:
            awaiting_components = False
            last_msg_idx = None
            label = m.group("label").upper()
            bits = m.group("bits")
            ttag = m.group("hms")
            if acc_label is None:
                acc_label = label
                acc_time = ttag
                acc_bits.append(bits)
                acc_components = []
                j = i
                lookahead_count = 0
                while lookahead_count < 2 and (j + 1) < n:
                    nxt = lines[j + 1].rstrip("\n")
                    nm = IXL_LINE_RE.search(nxt)
                    if nm:
                        nxt_label = nm.group("label").upper()
                        nxt_time = nm.group("hms")
                        if nxt_label == label and _within_point_one_seconds(ttag, nxt_time):
                            acc_bits.append(nm.group("bits"))
                            j += 1
                            lookahead_count += 1
                            continue
                    comp = _extract_component_from_line(nxt)
                    if comp:
                        comp_time = _time_from_any_line(nxt)
                        if comp_time and _within_point_one_seconds(ttag, comp_time):
                            acc_components.append(comp)
                            j += 1
                            continue
                    break
                i = j + 1
                continue
            else:
                if label == acc_label and _within_point_one_seconds(acc_time, ttag):
                    j = i
                    acc_bits.append(bits)
                    lookahead_count = 0
                    while lookahead_count < 2 and (j + 1) < n:
                        nxt = lines[j + 1].rstrip("\n")
                        nm = IXL_LINE_RE.search(nxt)
                        if nm:
                            nxt_label = nm.group("label").upper()
                            nxt_time = nm.group("hms")
                            if nxt_label == acc_label and _within_point_one_seconds(acc_time, nxt_time):
                                acc_bits.append(nm.group("bits"))
                                j += 1
                                lookahead_count += 1
                                continue
                        comp = _extract_component_from_line(nxt)
                        if comp:
                            comp_time = _time_from_any_line(nxt)
                            if comp_time and _within_point_one_seconds(acc_time, comp_time):
                                acc_components.append(comp)
                                j += 1
                                continue
                        break
                    i = j + 1
                    continue
                else:
                    idx = _flush_ixl_group(
                        acc_label, acc_bits, acc_time, messages, start_dt, end_dt, acc_components
                    )
                    if idx is not None and (
                        messages[idx]["msg_type"] == "Control (12 01)" or messages[idx]["msg_type"] == "Ind (12 8B)"
                    ):
                        awaiting_components = True
                        last_msg_idx = idx
                    else:
                        awaiting_components = False
                        last_msg_idx = None
                    acc_label, acc_bits, acc_time = label, [], ttag
                    acc_components = []
                    j = i
                    acc_bits.append(bits)
                    lookahead_count = 0
                    while lookahead_count < 2 and (j + 1) < n:
                        nxt = lines[j + 1].rstrip("\n")
                        nm = IXL_LINE_RE.search(nxt)
                        if nm:
                            nxt_label = nm.group("label").upper()
                            nxt_time = nm.group("hms")
                            if nxt_label == acc_label and _within_point_one_seconds(acc_time, nxt_time):
                                acc_bits.append(nm.group("bits"))
                                j += 1
                                lookahead_count += 1
                                continue
                        comp = _extract_component_from_line(nxt)
                        if comp:
                            comp_time = _time_from_any_line(nxt)
                            if comp_time and _within_point_one_seconds(acc_time, comp_time):
                                acc_components.append(comp)
                                j += 1
                                continue
                        break
                    i = j + 1
                    continue
        if IXL_END_EXECUTE_RE.search(line):
            idx = _flush_ixl_group(acc_label, acc_bits, acc_time, messages, start_dt, end_dt, acc_components)
            if idx is not None and (
                messages[idx]["msg_type"] == "Control (12 01)" or messages[idx]["msg_type"] == "Ind (12 8B)"
            ):
                awaiting_components = True
                last_msg_idx = idx
            else:
                awaiting_components = False
                last_msg_idx = None
            acc_label, acc_bits, acc_time = None, [], None
            acc_components = []
            i += 1
            continue
        if acc_bits:
            comp = _extract_component_from_line(line)
            if comp:
                comp_time = TIME_PATTERN.search(line)
                comp_time = comp_time.group(0) if comp_time else None
                if comp_time and acc_time:
                    try:
                        if abs((parse_time_flexible(comp_time) - parse_time_flexible(acc_time)).total_seconds()) <= 0.1:
                            acc_components.append(comp)
                            i += 1
                            continue
                    except Exception:
                        pass
            idx = _flush_ixl_group(acc_label, acc_bits, acc_time, messages, start_dt, end_dt, acc_components)
            if idx is not None and (
                messages[idx]["msg_type"] == "Control (12 01)" or messages[idx]["msg_type"] == "Ind (12 8B)"
            ):
                awaiting_components = True
                last_msg_idx = idx
            else:
                awaiting_components = False
                last_msg_idx = None
            acc_label, acc_bits, acc_time = None, [], None
            acc_components = []
        if awaiting_components and last_msg_idx is not None:
            comp = _extract_component_from_line(line)
            if comp:
                messages[last_msg_idx]["component_lines"].append(comp)
                i += 1
                continue
        i += 1

    if acc_bits:
        _flush_ixl_group(acc_label, acc_bits, acc_time, messages, start_dt, end_dt, acc_components)

    for msg in messages:
        lines_join = msg.get("component_lines", [])
        msg["component"] = "\r\n".join(lines_join) if lines_join else ""
        msg.pop("component_lines", None)
    return messages

# ========================
# Excel helpers
# ========================

def adjust_ixl_component_column(
    ws,
    header_row_idx: int = 2,
    start_row: int = 2,
    ixl_component_col_index: int = 6,
    min_width: int = 28,
    max_width: int = 60,
):
    """Wrap text & set readable width for IXL 'component' column (F by default)."""
    col_letter = get_column_letter(ixl_component_col_index)
    for row in ws.iter_rows(
        min_row=start_row, max_row=ws.max_row, min_col=ixl_component_col_index, max_col=ixl_component_col_index
    ):
        for cell in row:
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    longest_line_len = 0
    for row in ws.iter_rows(
        min_row=start_row, max_row=ws.max_row, min_col=ixl_component_col_index, max_col=ixl_component_col_index
    ):
        for cell in row:
            val = str(cell.value) if cell.value is not None else ""
            for ln in val.splitlines():
                longest_line_len = max(longest_line_len, len(ln))
    target_width = max(min_width, min(max_width, longest_line_len + 2))
    ws.column_dimensions[col_letter].width = target_width
    fixed_height = 15
    for r in range(start_row, ws.max_row + 1):
        ws.row_dimensions[r].height = fixed_height

# ========================
# IXL dataframe builder aligned to Wireshark rows
# ========================

def build_ixl_dataframe(
    ixl_file_path: str,
    ws_entries: list,
    start_time_str: str = None,
    end_time_str: str = None,
    max_diff_minutes: float = MAX_IXL_WS_DIFF_MINUTES,
) -> pd.DataFrame:
    row_count = len(ws_entries)
    aligned_empty = [
        {"time tag": "", "msg type": "", "data (hex)": "", "direction": "", "component": ""}
        for _ in range(row_count)
    ]
    if not ixl_file_path or not os.path.isfile(ixl_file_path):
        return pd.DataFrame(aligned_empty, columns=["time tag", "msg type", "data (hex)", "direction", "component"])

    start_dt = end_dt = None
    try:
        if start_time_str:
            start_dt = parse_time_flexible(start_time_str)
        if end_time_str:
            end_dt = parse_time_flexible(end_time_str)
    except Exception:
        start_dt = end_dt = None

    ixl_msgs = parse_ixl_file(ixl_file_path, start_dt=start_dt, end_dt=end_dt)
    unmatched_ixl = [
        {
            "time_tag": m["time_tag"],
            "msg_type": m["msg_type"],
            "component": m.get("component", ""),
            "data_pretty": m["data_hex"],
            "data_norm": normalize_hex_no_spaces(m["data_hex"]),
        }
        for m in ixl_msgs
    ]

    rows = [
        {"time tag": "", "msg type": "", "data (hex)": "", "direction": "", "component": ""}
        for _ in range(row_count)
    ]

    for i, ws_row in enumerate(ws_entries):
        if not ws_row or len(ws_row) < 4:
            continue
        ws_time_str = ws_row[0] or ""
        ws_msg_type = ws_row[1]
        ws_payload_pretty = ws_row[3]
        if not ws_msg_type or not ws_payload_pretty or not ws_time_str:
            continue
        if ws_msg_type not in ("Control (12 01)", "Ind (12 8B)"):
            continue
        ws_payload_norm = normalize_hex_no_spaces(ws_payload_pretty)
        try:
            ws_dt = parse_time_flexible(ws_time_str)
        except Exception:
            continue
        found_idx = None
        for j, im in enumerate(unmatched_ixl):
            if im["msg_type"] != ws_msg_type:
                continue
            if im["data_norm"] != ws_payload_norm:
                continue
            try:
                ixl_dt = parse_time_flexible(im["time_tag"])
            except Exception:
                continue
            if minutes_abs(ws_dt, ixl_dt) <= float(max_diff_minutes):
                found_idx = j
                break
        if found_idx is not None:
            im = unmatched_ixl.pop(found_idx)
            rows[i]["time tag"] = im["time_tag"]
            rows[i]["msg type"] = im["msg_type"]
            rows[i]["data (hex)"] = im["data_pretty"]
            component_text = im.get("component", "")
            if component_text:
                component_text = "See more ...\r\n" + component_text
            rows[i]["component"] = component_text

    return pd.DataFrame(rows, columns=["time tag", "msg type", "data (hex)", "direction", "component"])

# ========================
# Core analysis
# ========================

def analyze_logs(
    ixl_file_path,
    log_file_path,
    pcap_file_path,
    ixl_excel_file_path,
    start_time_str,
    end_time_str,
    packetswitch_file_path,
    target_address,
    filename_suffix,
):
    output_file_path = (
        f"log_packet_analysis_output_{filename_suffix or datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    )
    if os.path.exists(output_file_path) and is_file_open(output_file_path):
        messagebox.showerror("Error", "The output file is currently open. Please close it and try again.")
        return

    cm_entries = []  # Communication Manager
    ws_entries = []  # Wireshark
    time_differences = []
    last_two_sequences = []
    recent_sequence_times = {}

    if log_file_path != False:
        try:
            start_dt = parse_time_only(start_time_str)
            end_dt = parse_time_only(end_time_str)
        except Exception as e:
            messagebox.showerror("Error", f"Invalid time format: {e}")
            return
        with open(log_file_path, "r") as log_file:
            lines = log_file.readlines()
        packets = rdpcap(pcap_file_path)

        # Packetswitch read
        html_text = ""
        if packetswitch_file_path.lower().endswith(".html"):
            with open(packetswitch_file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                html_text = soup.get_text()
        elif packetswitch_file_path.lower().endswith(".docx"):
            if is_file_open(packetswitch_file_path):
                messagebox.showerror(
                    "Error", "The Packetswitch file is currently open. Please close it and try again."
                )
                return
            try:
                html_text = "\n".join([para.text for para in Document(packetswitch_file_path).paragraphs])
            except Exception:
                with open(packetswitch_file_path, "r", encoding="utf-8", errors="ignore") as f:
                    html_text = f.read()

        # Iterate CM log lines
        i = 0
        while i < len(lines):
            line = lines[i]
            time_match = TIME_PATTERN.search(line)
            if time_match:
                last_timestamp = time_match.group()
                try:
                    log_dt = parse_time_only(last_timestamp)
                except Exception:
                    i += 1
                    continue
                if not (start_dt <= log_dt <= end_dt):
                    i += 1
                    continue
            else:
                i += 1
                continue

            if "02 02" in line:
                try:
                    index = line.index("02 02")
                    after = line[index + len("02 02") :].strip()
                    after_parts = after.split()
                    if len(after_parts) >= 3 and after_parts[0] == "02":
                        msg_number = "02"
                        msg_type_raw = f"{after_parts[1]} {after_parts[2]}"
                    else:
                        pre_match = re.search(r"(\S{2})\s+02\s+02", line)
                        post_match = re.search(r"02\s+02\s+(\S{2})\s+(\S{2})", line)
                        msg_number = pre_match.group(1) if pre_match else None
                        msg_type_raw = (
                            f"{post_match.group(1)} {post_match.group(2)}" if post_match else None
                        )
                    line_type = "BASIC" if "BASIC" in line else "INFO" if "INFO" in line else ""
                    msg_type = msg_type_raw
                    if msg_type_raw == "12 8B":
                        msg_type = "Ind (12 8B)"
                    elif msg_type_raw == "12 01":
                        msg_type = "Control (12 01)"
                    elif msg_type_raw == "12 48":
                        msg_type = "Recall (12 48)"
                    hex_sequence = f"{msg_number} 02 02 {msg_type_raw}" if msg_number and msg_type_raw else None
                except Exception:
                    i += 1
                    continue

                if not hex_sequence or hex_sequence in last_two_sequences:
                    i += 1
                    continue
                if hex_sequence in recent_sequence_times:
                    if (log_dt - recent_sequence_times[hex_sequence]) <= timedelta(seconds=30):
                        i += 1
                        continue
                try:
                    search_bytes = bytes.fromhex(hex_sequence)
                except Exception:
                    i += 1
                    continue

                pcap_ts_str, time_diff_str, _, ws_hex_data = find_pcap_time(
                    search_bytes, log_dt, packets, target_address, pcap_file_path, msg_type_raw
                )
                if time_diff_str != "Found but overflow":
                    hex_data = extract_hex_data(lines, i, msg_type_raw, line_type)
                    cm_entries.append([last_timestamp, msg_type, msg_number, hex_data, ""])
                    ws_entries.append([pcap_ts_str or "", "" if pcap_ts_str is None else msg_type, "" if pcap_ts_str is None else msg_number, ws_hex_data, ""])
                    time_differences.append([time_diff_str])
                    last_two_sequences.append(hex_sequence)
                    if len(last_two_sequences) > 2:
                        last_two_sequences.pop(0)
                    recent_sequence_times[hex_sequence] = log_dt
                i += 1
            else:
                i += 1

    # SECOND ITERATION: Wireshark-only scan for 02 02 and RF_ACK when CM log not provided
    last_rf_ack_time = None
    if log_file_path == False:
        found_cm_times = set()
    else:
        found_cm_times = set(entry[0] for entry in cm_entries)

    additional_ws_entries = []
    additional_cm_entries = []
    additional_time_differences = []

    if log_file_path == False:
        packets = rdpcap(pcap_file_path)
        try:
            start_dt = parse_time_only(start_time_str)
            end_dt = parse_time_only(end_time_str)
        except Exception as e:
            messagebox.showerror("Error", f"Invalid time format: {e}")
            return
        filtered_packets = []
        for pkt in packets:
            try:
                pkt_time = datetime.fromtimestamp(float(pkt.time))
                pkt_time_only = datetime.strptime(pkt_time.strftime("%H:%M:%S.%f")[:-3], "%H:%M:%S.%f")
                if start_dt <= pkt_time_only <= end_dt:
                    filtered_packets.append(pkt)
            except Exception:
                continue
        packets = filtered_packets

        html_text = ""
        if packetswitch_file_path.lower().endswith(".html"):
            with open(packetswitch_file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                html_text = soup.get_text()
        elif packetswitch_file_path.lower().endswith(".docx"):
            if is_file_open(packetswitch_file_path):
                messagebox.showerror("Error", "The Packetswitch file is currently open. Please close it and try again.")
                return
            try:
                html_text = "\n".join([para.text for para in Document(packetswitch_file_path).paragraphs])
            except Exception:
                with open(packetswitch_file_path, "r", encoding="utf-8", errors="ignore") as f:
                    html_text = f.read()

        for pkt in packets:
            if Raw in pkt:
                data_bytes = pkt[Raw].load
                if is_valid_packet(pkt, data_bytes, target_address):
                    hex_str = data_bytes.hex().upper()
                    hex_parts = [hex_str[i : i + 2] for i in range(0, len(hex_str), 2)]
                    for i in range(len(hex_parts) - 3):
                        # Detect RF_ACK (8C .... 34/38)
                        for j in range(len(hex_parts) - 4):
                            if hex_parts[j] == "8C":
                                fourth_pair = hex_parts[j + 4]
                                if fourth_pair in ["34", "38"]:
                                    rf_ack_time = datetime.fromtimestamp(float(pkt.time))
                                    if last_rf_ack_time and (rf_ack_time - last_rf_ack_time) <= timedelta(seconds=5):
                                        continue
                                    msg_type = "RF_ACK (Inbound)" if fourth_pair == "34" else "RF_ACK (Outbound)"
                                    last_rf_ack_time = rf_ack_time
                                    pcap_time = rf_ack_time.strftime("%H:%M:%S.%f")[:-3]
                                    additional_cm_entries.append(["Not found", "", "", "", ""])
                                    additional_ws_entries.append([pcap_time, msg_type, "", "", ""])
                                    additional_time_differences.append([""])
                                    break
                        # Look for 02 02
                        if hex_parts[i] == "02" and hex_parts[i + 1] == "02":
                            suffix = target_address.hex()[-2:]
                            search_window = hex_parts[max(0, i - 11) : i]
                            if suffix not in search_window:
                                continue
                            if i > 0 and hex_parts[i + 2] == "02":
                                msg_number = "02"
                                msg_type_raw = f"{hex_parts[i + 3]} {hex_parts[i + 4]}"
                            else:
                                msg_number = hex_parts[i - 1] if i > 0 else None
                                msg_type_raw = f"{hex_parts[i + 2]} {hex_parts[i + 3]}"
                            msg_type = msg_type_raw
                            if msg_type_raw == "04 D0":
                                continue
                            if msg_type_raw == "12 8B":
                                msg_type = "Ind (12 8B)"
                            elif msg_type_raw == "12 01":
                                msg_type = "Control (12 01)"
                            elif msg_type_raw == "12 48":
                                msg_type = "Recall (12 48)"
                            hex_sequence = f"{msg_number} 02 02 {msg_type_raw}" if msg_number and msg_type_raw else None
                            if log_file_path != False:
                                if hex_sequence in [
                                    f"{entry[2]} 02 02 {entry[1].split('(')[-1].strip(')')}" for entry in cm_entries if entry[0] != "Not found"
                                ]:
                                    continue
                            pcap_time = datetime.fromtimestamp(float(pkt.time)).strftime("%H:%M:%S.%f")[:-3]
                            if not hex_sequence or hex_sequence in last_two_sequences:
                                continue
                            if hex_sequence in recent_sequence_times:
                                if (datetime.fromtimestamp(float(pkt.time)) - recent_sequence_times[hex_sequence]) <= timedelta(seconds=30):
                                    continue
                            last_two_sequences.append(hex_sequence)
                            if len(last_two_sequences) > 2:
                                last_two_sequences.pop(0)
                            recent_sequence_times[hex_sequence] = datetime.fromtimestamp(float(pkt.time))
                            if pcap_time not in found_cm_times:
                                additional_cm_entries.append(["Not found", "", "", "", ""])
                            fallback_hex = " ".join([data_bytes.hex()[k : k + 2] for k in range(0, len(data_bytes.hex()), 2)])
                            ws_hex_data = extract_ws_hex_data(pcap_file_path, pcap_time, msg_type_raw, fallback_hex)
                            additional_ws_entries.append([pcap_time, msg_type, msg_number, ws_hex_data, ""])
                            additional_time_differences.append([""])

    cm_entries.extend(additional_cm_entries)
    ws_entries.extend(additional_ws_entries)
    time_differences.extend(additional_time_differences)

    # Filter by user time range (either CM or WS time inside the window)
    if log_file_path != False:
        filtered_entries = []
        for i, entry in enumerate(cm_entries):
            cm_time_str = entry[0]
            ws_time_str = ws_entries[i][0]
            try:
                cm_time = parse_time_only(cm_time_str) if cm_time_str != "Not found" else None
                ws_time = parse_time_only(ws_time_str) if ws_time_str else None
                if (cm_time and start_dt <= cm_time <= end_dt) or (ws_time and start_dt <= ws_time <= end_dt):
                    filtered_entries.append((entry, ws_entries[i], time_differences[i]))
            except Exception:
                continue
        cm_entries = [e[0] for e in filtered_entries]
        ws_entries = [e[1] for e in filtered_entries]
        time_differences = [e[2] for e in filtered_entries]

    # Sort: valid CM time-tagged entries, insert 'Not found' by WS time
    combined_entries = list(zip(cm_entries, ws_entries, time_differences))
    valid_entries = [e for e in combined_entries if e[0][0] != "Not found"]
    valid_entries.sort(key=lambda x: parse_time_only(x[0][0]))
    did_not_find_entries = [e for e in combined_entries if e[0][0] == "Not found"]
    for entry in did_not_find_entries:
        ws_time_str = entry[1][0]
        try:
            ws_time = parse_time_only(ws_time_str)
        except Exception:
            ws_time = datetime.max
        inserted = False
        for i, valid_entry in enumerate(valid_entries):
            try:
                next_ws_time = parse_time_only(valid_entry[1][0])
                if ws_time < next_ws_time:
                    valid_entries.insert(i, entry)
                    inserted = True
                    break
            except Exception:
                continue
        if not inserted:
            valid_entries.append(entry)

    cm_entries = [e[0] for e in valid_entries]
    ws_entries = [e[1] for e in valid_entries]
    time_differences = [e[2] for e in valid_entries]

    for entry in ws_entries:
        msg_type_raw = entry[1].split("(")[-1].strip(")") if "(" in entry[1] else entry[1]
        entry[3] = process_ws_hex_data(entry[3], msg_type_raw)

    # Build DataFrames
    row_count = len(cm_entries)
    ixl_df = build_ixl_dataframe(ixl_file_path, ws_entries, start_time_str, end_time_str)
    diff2_df = pd.DataFrame({" ": [""] * row_count})
    cm_df = pd.DataFrame(cm_entries, columns=["time tag", "msg type", "msg number", "data (hex)", "direction"])
    ws_df = pd.DataFrame(ws_entries, columns=["time tag", "msg type", "msg number", "data (hex)", "UDP/TCP"])
    diff_df = pd.DataFrame(time_differences, columns=[" "])

    # Packetswitch integration
    html_text = ""
    if packetswitch_file_path and os.path.isfile(packetswitch_file_path):
        if packetswitch_file_path.lower().endswith(".html"):
            with open(packetswitch_file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                html_text = soup.get_text()
        elif packetswitch_file_path.lower().endswith(".docx"):
            if is_file_open(packetswitch_file_path):
                messagebox.showerror("Error", "The Packetswitch file is currently open. Please close it and try again.")
                return
            try:
                html_text = "\n".join([para.text for para in Document(packetswitch_file_path).paragraphs])
            except Exception:
                with open(packetswitch_file_path, "r", encoding="utf-8", errors="ignore") as f:
                    html_text = f.read()

    if not html_text.strip():
        packetswitch_df = pd.DataFrame(
            {
                "time tag": [""] * len(ws_entries),
                "component": [""] * len(ws_entries),
                "data (hex)": [""] * len(ws_entries),
                "data type": [""] * len(ws_entries),
            }
        )
    else:
        is_generic_report = "Generic Report Results" in html_text
        if not is_generic_report:
            packetswitch_times = []
            packetswitch_codes = []
            packetswitch_count = {}
            for entry in ws_entries:
                msg_type_raw = entry[1].split("(")[-1].strip(")") if "(" in entry[1] else entry[1]
                if msg_type_raw not in ["12 8B", "12 01", "12 48"]:
                    packetswitch_times.append("")
                    packetswitch_codes.append("")
                    continue
                pcap_time = entry[0]
                msg_type = entry[1]
                code = ""
                if pcap_time:
                    time_tag = pcap_time.split(".")[0]
                    key = (time_tag, msg_type)
                    count = packetswitch_count.get(key, 0)
                    idx = -1
                    for _ in range(count + 1):
                        idx = html_text.find(time_tag, idx + 1)
                        if idx == -1:
                            break
                    if idx != -1:
                        after = html_text[idx:]
                        underscore_idx = after.find("_")
                        if underscore_idx != -1 and len(after) > underscore_idx + 2:
                            raw_code = after[underscore_idx + 1 : underscore_idx + 3]
                            if raw_code == "CR" and msg_type == "Control (12 01)":
                                code = "Control"
                            elif raw_code == "IR" and msg_type == "Ind (12 8B)":
                                code = "Ind"
                            elif raw_code == "R_" and msg_type == "Recall (12 48)":
                                code = "Recall"
                            if code:
                                packetswitch_times.append(time_tag)
                                packetswitch_codes.append(code)
                                packetswitch_count[key] = count + 1
                                continue
                # tenth-digit 9 retry (+1 second)
                tenth_digit_is_nine = False
                try:
                    fractional = pcap_time.split(".")[-1]
                    if len(fractional) >= 1 and fractional[0] == "9":
                        tenth_digit_is_nine = True
                except Exception:
                    pass
                if not code and tenth_digit_is_nine:
                    adjusted_time = (datetime.strptime(time_tag, "%H:%M:%S") + timedelta(seconds=1)).strftime("%H:%M:%S")
                    key = (adjusted_time, msg_type)
                    count = packetswitch_count.get(key, 0)
                    idx = -1
                    for _ in range(count + 1):
                        idx = html_text.find(adjusted_time, idx + 1)
                        if idx == -1:
                            break
                    if idx != -1:
                        after = html_text[idx:]
                        underscore_idx = after.find("_")
                        if underscore_idx != -1 and len(after) > underscore_idx + 2:
                            raw_code = after[underscore_idx + 1 : underscore_idx + 3]
                            if raw_code == "CR" and msg_type == "Control (12 01)":
                                code = "Control"
                            elif raw_code == "IR" and msg_type == "Ind (12 8B)":
                                code = "Ind"
                            elif raw_code == "R_" and msg_type == "Recall (12 48)":
                                code = "Recall"
                            if code:
                                packetswitch_times.append(adjusted_time)
                                packetswitch_codes.append(code)
                                packetswitch_count[key] = count + 1
                                continue
                packetswitch_times.append("")
                packetswitch_codes.append("")
            packetswitch_df = pd.DataFrame(
                {"time tag": packetswitch_times, "component": [""] * len(packetswitch_times), "data (hex)": [""] * len(packetswitch_times), "data type": packetswitch_codes}
            )
        else:
            packetswitch_times = []
            packetswitch_components = []
            packetswitch_codes = []
            packetswitch_data = []
            ps_lines = html_text.strip().split("\n")
            parsed_ps_lines = []
            for line in ps_lines:
                match = re.match(r".*?(\d{2}:\d{2}:\d{2})\s+([A-Za-z()]+)\s.*?:\s*([0-9A-Fa-f ]+)$", line)
                if match:
                    time_tag = match.group(1)
                    description = match.group(2)
                    data = match.group(3).strip()
                    parsed_ps_lines.append((time_tag, description, data))
            for ws_entry in ws_entries:
                ws_time_str = ws_entry[0].split(".")[0]
                try:
                    ws_time = datetime.strptime(ws_time_str, "%H:%M:%S")
                except Exception:
                    ws_time = None
                ws_data = ws_entry[3].replace(" ", "").upper()
                found_match = False
                for time_tag, description, ps_data in parsed_ps_lines:
                    ps_data_clean = ps_data.replace(" ", "").upper()
                    if ps_data_clean.startswith(ws_data):
                        try:
                            ps_time = datetime.strptime(time_tag, "%H:%M:%S")
                            time_diff = abs((ps_time - ws_time).total_seconds()) if ws_time else None
                        except Exception:
                            time_diff = None
                        if time_diff is not None and time_diff <= 1:
                            if ws_entry[1] in ("RF_ACK (Outbound)", "RF_ACK (Inbound)"):
                                continue
                            if description == "Indic(RF)":
                                if ws_entry[1] != "Ind (12 8B)":
                                    continue
                                else:
                                    description = "Ind"
                            if description == "Rf":
                                if ws_entry[1] != "Control (12 01)":
                                    continue
                                else:
                                    description = "Control"
                            if description == "Indicate":
                                continue
                            packetswitch_times.append(time_tag)
                            packetswitch_components.append("")
                            packetswitch_codes.append(description)
                            packetswitch_data.append(ps_data)
                            found_match = True
                            break
                if not found_match:
                    packetswitch_times.append("")
                    packetswitch_components.append("")
                    packetswitch_codes.append("")
                    packetswitch_data.append("")
            packetswitch_df = pd.DataFrame(
                {
                    "time tag": packetswitch_times,
                    "component": packetswitch_components,
                    "data (hex)": packetswitch_data,
                    "data type": packetswitch_codes,
                }
            )

    combined_df = pd.concat(
        [
            pd.DataFrame({"Number": list(range(1, len(cm_df) + 1))}),
            ixl_df,
            diff2_df,
            cm_df,
            diff_df,
            ws_df,
            packetswitch_df,
        ],
        axis=1,
    )

    header_row = [
        " ",
        "IXL Log",
        "",
        "",
        "",
        "",
        "Time Difference",
        "Communication Manager Log",
        "",
        "",
        "",
        "",
        "Time Difference",
        "Wireshark Log",
        "",
        "",
        "",
        "",
        "Packetswitch Data",
        "",
        "",
        "",
    ]

    with pd.ExcelWriter(output_file_path, engine="openpyxl") as writer:
        pd.DataFrame([header_row]).to_excel(writer, index=False, header=False)
        combined_df.to_excel(writer, index=False, startrow=1)

    wb = load_workbook(output_file_path)
    ws = wb.active
    ws.freeze_panes = "A3"
    ws.merge_cells(start_row=1, start_column=2, end_row=1, end_column=6)
    ws.merge_cells(start_row=1, start_column=7, end_row=1, end_column=7)
    ws.merge_cells(start_row=1, start_column=8, end_row=1, end_column=12)
    ws.merge_cells(start_row=1, start_column=13, end_row=1, end_column=13)
    ws.merge_cells(start_row=1, start_column=14, end_row=1, end_column=18)
    ws.merge_cells(start_row=1, start_column=19, end_row=1, end_column=22)

    for col in [2, 7, 8, 13, 14, 19]:
        cell = ws.cell(row=1, column=col)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=22):
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center")
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=6, max_col=6):
        for cell in row:
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    for col in ws.iter_cols(min_row=2):
        max_length = 0
        column = col[0].column
        column_letter = get_column_letter(column)
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column_letter].width = max_length + 2
    adjust_ixl_component_column(ws, header_row_idx=2, start_row=2, ixl_component_col_index=6, min_width=28, max_width=60)

    yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    COL_WS_TYPE = 15
    COL_WS_NUM = 16
    last_by_group = {}
    for r in range(2, ws.max_row + 1):
        ws_type_val = ws.cell(row=r, column=COL_WS_TYPE).value
        ws_num_val = ws.cell(row=r, column=COL_WS_NUM).value
        if not ws_type_val or not ws_num_val:
            continue
        raw_type = ws_type_val.split("(")[-1].strip(")") if "(" in ws_type_val else ws_type_val
        if raw_type.upper() == "08 42":
            continue
        group = _group_from_msg_type_cell(ws_type_val)
        try:
            current = int(str(ws_num_val).strip(), 16)
        except Exception:
            continue
        last = last_by_group.get(group)
        if last is None:
            last_by_group[group] = current
            continue
        expected = (last + 2) % 256
        if current != expected:
            ws.cell(row=r, column=COL_WS_NUM).fill = yellow_fill
        last_by_group[group] = current

    wb.save(output_file_path)
    messagebox.showinfo(
        "Success",
        "Output has been saved to: log_packet_analysis_output.xlsx\nThe Excel file is ready to be viewed.",
    )

# ========================
# GUI helpers
# ========================

def browse_ixl_file():
    filename = filedialog.askopenfilename(title="Select IXL Log File", filetypes=[("Text Files", "*.txt")])
    ixl_file_entry.delete(0, tk.END)
    ixl_file_entry.insert(0, filename)


def browse_log_file():
    filename = filedialog.askopenfilename(
        title="Select Communication Manager Log File", filetypes=[("Text Files", "*.txt")]
    )
    log_file_entry.delete(0, tk.END)
    log_file_entry.insert(0, filename)


def browse_pcap_file():
    filename = filedialog.askopenfilename(title="Select Wireshark PCAP File", filetypes=[("PCAP Files", "*.pcap")])
    pcap_file_entry.delete(0, tk.END)
    pcap_file_entry.insert(0, filename)


def browse_packetswitch_file():
    filename = filedialog.askopenfilename(
        title="Select Packetswitch File", filetypes=[("HTML or DOCX Files", "*.html *.docx")]
    )
    packetswitch_file_entry.delete(0, tk.END)
    packetswitch_file_entry.insert(0, filename)


def browse_ixl_excel_file():
    filename = filedialog.askopenfilename(title="Select IXL Excel File", filetypes=[("Excel Files", "*.xlsx")])
    ixl_excel_file_entry.delete(0, tk.END)
    ixl_excel_file_entry.insert(0, filename)


def run_analysis():
    ixl_file = ixl_file_entry.get()
    log_file = log_file_entry.get()
    pcap_file = pcap_file_entry.get()
    packetswitch_file = packetswitch_file_entry.get()
    ixl_excel_file = ixl_excel_file_entry.get()
    start_time = start_time_entry.get()
    end_time = end_time_entry.get()
    target_address_hex = target_address_entry.get().replace("0", "a").replace(".", "")

    import string

    def is_valid_hex(s):
        return len(s) % 2 == 0 and all(c in string.hexdigits for c in s)

    if not is_valid_hex(target_address_hex):
        messagebox.showerror("Error", "Target address must be even-length and contain only hex characters.")
        return
    if not os.path.isfile(pcap_file):
        messagebox.showerror("Error", "Please select a valid PCAP file path.")
        return

    if not ixl_file:
        ixl_file = False
    if not log_file:
        log_file = False
    if not ixl_excel_file:
        ixl_excel_file = False

    analyze_logs(
        ixl_file,
        log_file,
        pcap_file,
        ixl_excel_file,
        start_time,
        end_time,
        packetswitch_file,
        bytes.fromhex(target_address_hex),
        filename_suffix_entry.get(),
    )

# ========================
# UI layout
# ========================
root = tk.Tk()
root.title("Log Packet Analysis Tool")

# Row 0
tk.Label(root, text="IXL Text File:").grid(row=0, column=0, sticky="w", padx=10, pady=5)
ixl_file_entry = tk.Entry(root, width=60)
ixl_file_entry.grid(row=0, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_ixl_file).grid(row=0, column=2, padx=10)

# Row 1
tk.Label(root, text="Communication Manager Log Text File:").grid(row=1, column=0, sticky="w", padx=10, pady=5)
log_file_entry = tk.Entry(root, width=60)
log_file_entry.grid(row=1, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_log_file).grid(row=1, column=2, padx=10)

# Row 2
tk.Label(root, text="Wireshark PCAP File:").grid(row=2, column=0, sticky="w", padx=10, pady=5)
pcap_file_entry = tk.Entry(root, width=60)
pcap_file_entry.grid(row=2, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_pcap_file).grid(row=2, column=2, padx=10)

# Row 3
tk.Label(root, text="Packetswitch Data File:").grid(row=3, column=0, sticky="w", padx=10, pady=5)
packetswitch_file_entry = tk.Entry(root, width=60)
packetswitch_file_entry.grid(row=3, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_packetswitch_file).grid(row=3, column=2, padx=10)

# Row 4
tk.Label(root, text="IXL Excel File:").grid(row=4, column=0, sticky="w", padx=10, pady=5)
ixl_excel_file_entry = tk.Entry(root, width=60)
ixl_excel_file_entry.grid(row=4, column=1, padx=10)
tk.Button(root, text="Browse", command=browse_ixl_excel_file).grid(row=4, column=2, padx=10)

# Row 5
tk.Label(root, text="Start Time (HH:MM:SS.sss):").grid(row=5, column=0, sticky="w", padx=10, pady=5)
start_time_entry = tk.Entry(root, width=20)
start_time_entry.grid(row=5, column=1, sticky="w", padx=10)

# Row 6
tk.Label(root, text="End Time (HH:MM:SS.sss):").grid(row=6, column=0, sticky="w", padx=10, pady=5)
end_time_entry = tk.Entry(root, width=20)
end_time_entry.grid(row=6, column=1, sticky="w", padx=10)

# Row 7
tk.Label(root, text="Target Address:").grid(row=7, column=0, sticky="w", padx=10, pady=5)
target_address_entry = tk.Entry(root, width=20)
target_address_entry.grid(row=7, column=1, sticky="w", padx=10)

# Row 8
tk.Label(root, text="Output suffix after 'log_packet_analysis_output':").grid(row=8, column=0, sticky="w", padx=10, pady=5)
filename_suffix_entry = tk.Entry(root, width=20)
filename_suffix_entry.grid(row=8, column=1, sticky="w", padx=10)

# Row 9
tk.Button(root, text="Run Analysis", command=run_analysis, bg="green", fg="white").grid(row=9, column=1, pady=20)

# Row 10
tk.Label(root, text="You will receive a message that the Excel file is ready to be viewed. This may take a few minutes", fg="blue").grid(row=10, column=0, columnspan=3, pady=10)

root.mainloop()
