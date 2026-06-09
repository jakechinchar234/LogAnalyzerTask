# Latest with Repeat time tag case: 

# analyzer_backend.py
"""
Backend logic for the Log Packet Analysis Tool.
This file maintains the original analysis logic and helpers, without the Tkinter UI.
"""

"""
if not pcap_file_path:
    print("No Wireshark file selected — running in CM-only mode")
"""

# ========================
# Imports
# ========================
import os
import re
import json
import subprocess
from datetime import datetime, timedelta

# GUI notifications used by backend (kept for parity with original behavior)
from tkinter import messagebox

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document

from scapy.all import rdpcap, Raw
from scapy.layers.inet import TCP

from openpyxl import load_workbook
from openpyxl.styles import Alignment, PatternFill
from openpyxl.utils import get_column_letter

from BitSwitchingDetector import build_component_map

def _normalize_atcs(addr: str) -> str:
    """Normalize ATCS address for comparison."""
    return "".join(c for c in addr if c.isalnum()).lower()

def load_location_excel_mapping(filepath):
    """
    Reads an Excel file with:
        Column A: ATCS Address (dotted)
        Column B: Location Name

    Returns:
        dict -> { normalized_address: location_name }
    """
    if not filepath or not os.path.isfile(filepath):
        return {}

    try:
        df = pd.read_excel(filepath)

        # Assume first two columns
        addr_col = df.columns[0]
        name_col = df.columns[1]

        mapping = {}

        for _, row in df.iterrows():
            addr = str(row[addr_col]).strip()
            name = str(row[name_col]).strip()

            if not addr or not name:
                continue

            norm = _normalize_atcs(addr)
            mapping[norm] = name

        return mapping

    except Exception as e:
        print("Excel load error:", e)
        return {}

def resolve_user_input_to_hex(addr_text: str, excel_map: dict = None) -> str:
    """
    Accepts:
        - Location Name (from Excel)
        - Dotted ATCS address
        - Raw hex

    Returns:
        hex string (first 10 chars ONLY)
    """
    if not addr_text:
        return ""

    addr_text = addr_text.strip()

    # Case 1: Name looked up in Excel
    if excel_map:
        for norm, name in excel_map.items():
            if name == addr_text:
                # convert normalized back to dotted then to hex
                dotted = f"{norm[0]}.{norm[1:4]}.{norm[4:7]}.{norm[7:10]}"
                raw = dotted.replace(".", "").replace("0", "a")
                return raw[:10]

    # Case 2: Dotted
    if "." in addr_text:
        raw = addr_text.replace(".", "").replace("0", "a")
        return raw[:10]

    # Case 3: already hex
    raw = re.sub(r'[^0-9A-Fa-f]', '', addr_text)
    return raw[:10]

# ========================
# Constants & Regex
# ========================
TIME_PATTERN = re.compile(r"\b\d{2}:\d{2}:\d{2}\.\d{2}\b")

IXL_LABEL_TO_WS_TYPE = {
    "CTL": "Control (12 01)",
    "IND": "Ind (12 8B)",
}
IXL_LINE_RE = re.compile(
    r"^(?:\w{3}\s+)?"  # Optional weekday (e.g., Wed)
    r"(?P<md>\d{2}-\d{2}(?:-\d{4})?)\s+"  # MM-DD or MM-DD-YYYY
    r"(?P<hms>\d{2}:\d{2}:\d{2}\.\d{2}).*?\b"
    r"(?P<label>CTL|IND)\(\d{1,3}-\d{1,3}\):\s+(?P<bits>[01]{8})"
)
IXL_END_EXECUTE_RE = re.compile(r"Execute issued", re.IGNORECASE)
MAX_IXL_WS_DIFF_MINUTES = 8

# ========================
# Small utilities
# ========================

def parse_time_only(ts_str: str) -> datetime:
    if "." not in ts_str:
        ts_str += ".0"
    return datetime.strptime(ts_str, "%H:%M:%S.%f")


def parse_time_flexible(ts_str: str) -> datetime:
    if not ts_str:
        raise ValueError("Empty time string")
    if "." in ts_str:
        hms, frac = ts_str.split(".", 1)
        if len(frac) == 2:
            ts_norm = f"{hms}.{frac}0"
        elif len(frac) >= 3:
            ts_norm = f"{hms}.{frac[:3]}"
        else:
            ts_norm = f"{hms}.{frac.ljust(3, '0')}"
    else:
        ts_norm = ts_str + ".000"
    return datetime.strptime(ts_norm, "%H:%M:%S.%f")


def format_timedelta(td: timedelta) -> str:
    total_seconds = int(td.total_seconds())
    milliseconds = int(td.microseconds / 10000)
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours:02}:{minutes:02}:{seconds:02}.{milliseconds:02}"


def minutes_abs(dt_a: datetime, dt_b: datetime) -> float:
    a = dt_a.time(); b = dt_b.time()
    to_secs = lambda t: t.hour * 3600 + t.minute * 60 + t.second + t.microsecond / 1_000_000
    return abs(to_secs(a) - to_secs(b)) / 60.0


def is_file_open(filepath: str) -> bool:
    try:
        os.rename(filepath, filepath)
        return False
    except Exception:
        return True

# ========================
# Wireshark helpers
# ========================

def is_valid_packet(pkt, data_bytes: bytes, target_address: bytes) -> bool:
    return target_address in data_bytes and not TCP in pkt


def extract_ws_hex_data(pcap_path: str, ws_time_tag: str, msg_type_raw: str, fallback_hex: str = "") -> str:
    if msg_type_raw == "12 8B":
        wanted_label = "indication_bits"
    elif msg_type_raw == "12 01":
        wanted_label = "control_bits"
    else:
        return ""

    try:
        second_str = ws_time_tag.split(".")[0]
        display_filter = f'frame.time contains "{second_str}"'
        proc = subprocess.run(
            ["tshark", "-r", pcap_path, "-Y", display_filter, "-T", "json"],
            capture_output=True,
            text=True,
            check=True,
        )
        frames = json.loads(proc.stdout)
        for fr in frames:
            layers = fr.get("_source", {}).get("layers", {})
            t_epoch_list = layers.get("frame", {}).get("frame.time_epoch", [])
            if not t_epoch_list:
                continue
            try:
                t_epoch = float(t_epoch_list[0])
                candidate = datetime.fromtimestamp(t_epoch).strftime("%H:%M:%S.%f")[:-3]
            except Exception:
                continue
            if candidate != ws_time_tag:
                continue

            def find_val(node):
                if isinstance(node, dict):
                    if "showname" in node and isinstance(node["showname"], str):
                        prefix = f"{wanted_label}: "
                        if node["showname"].startswith(prefix):
                            return node["showname"][len(prefix):].strip()
                    for v in node.values():
                        val = find_val(v)
                        if val:
                            return val
                elif isinstance(node, list):
                    for item in node:
                        val = find_val(item)
                        if val:
                            return val
                return None

            val = find_val(layers)
            if val:
                raw = re.sub(r"[^0-9A-Fa-f]", "", val)
                return " ".join(raw[i:i+2].upper() for i in range(0, len(raw), 2))
    except Exception:
        pass
    return fallback_hex


def find_pcap_time(search_bytes: bytes, log_dt: datetime, packets, target_address: bytes, pcap_file_path: str, msg_type_raw: str):
    for pkt in packets:
        if Raw in pkt and target_address in pkt[Raw].load:
            data_bytes = pkt[Raw].load
            if search_bytes in data_bytes and target_address in data_bytes and not TCP in pkt:
                pcap_dt = datetime.fromtimestamp(float(pkt.time))
                pcap_ts_str = pcap_dt.strftime("%H:%M:%S.%f")[:-3]
                try:
                    pcap_time_only = datetime.strptime(pcap_ts_str, "%H:%M:%S.%f")
                    time_diff = pcap_time_only - log_dt
                    minutes_diff = time_diff.total_seconds() / 60
                    if pcap_time_only < (log_dt - timedelta(minutes=8)):
                        continue
                    if -8 <= minutes_diff <= 8:
                        fallback_hex = " ".join([
                            data_bytes.hex()[i:i+2] for i in range(0, len(data_bytes.hex()), 2)
                        ])
                        ws_hex_data = extract_ws_hex_data(pcap_file_path, pcap_ts_str, msg_type_raw, fallback_hex)
                        return pcap_ts_str, format_timedelta(time_diff), True, ws_hex_data
                    else:
                        return pcap_ts_str, "Found but overflow", False, ""
                except Exception:
                    continue
    return None, "LOST", False, ""

# ========================
# Hex processing
# ========================

def _normalize_group_from_raw(msg_type_raw: str) -> str:
    s = (msg_type_raw or "").upper().strip()
    if s in ("12 8B", "78 8B"):
        return "IND"
    if s in ("12 01", "12 48"):
        return "CTL_RECALL"
    if s in ("08 82", "08 83"):
        return "GROUP_0882_0883"
    return s


def _group_from_msg_type_cell(cell_val: str) -> str:
    val = (cell_val or "").strip()
    if "(" in val and ")" in val:
        raw = val.split("(")[-1].strip(")")
    else:
        raw = val
    return _normalize_group_from_raw(raw)


def process_ws_hex_data(ws_hex_data: str, msg_type_raw: str) -> str:
    if not ws_hex_data:
        return ws_hex_data
    parts = ws_hex_data.strip().split()
    idx = None
    for i in range(len(parts) - 3):
        if parts[i:i+4] == ["02","02","12","8b"] or parts[i:i+4] == ["02","02","12","01"]:
            idx = i + 4
            break
    if idx is None:
        return ws_hex_data
    trimmed = parts[idx:]
    trimmed = trimmed[4:] if len(trimmed) > 4 else []
    if msg_type_raw == "12 8B" and len(trimmed) > 0:
        trimmed = trimmed[:-1]
    return " ".join(trimmed)


def extract_hex_data(lines, start_index: int, msg_type_raw: str, line_type: str) -> str:
    collected = []
    found_sequence = False
    sequence_patterns = [["02","02","12","8B"],["02","02","12","01"]]
    i = start_index
    while i < len(lines):
        tokens = lines[i].strip().split()
        joined = " ".join(tokens).upper()
        if " TX" in joined or " RX" in joined or "TX:" in joined or "RX:" in joined:
            break
        if not found_sequence:
            for j in range(len(tokens)-3):
                if tokens[j:j+4] in sequence_patterns:
                    found_sequence = True
                    collected.extend([t for t in tokens[j+4:] if re.fullmatch(r"[0-9A-Fa-f]{2}", t)])
                    break
        else:
            collected.extend([t for t in tokens if re.fullmatch(r"[0-9A-Fa-f]{2}", t)])
        i += 1
    if not found_sequence or len(collected) < 3:
        return ""
    collected = collected[2:]
    length_hex = collected[0]
    length_bin = bin(int(length_hex, 16))[2:].zfill(8)
    length_bin_trimmed = length_bin[2:]
    length_val = int(length_bin_trimmed, 2)
    collected = collected[2:]
    if len(collected) > length_val:
        collected = collected[:length_val]
    return " ".join(collected)

# ========================
# IXL parsing & matching
# ========================

def reverse_bits_to_hex(bits_8: str) -> str:
    rev = bits_8[::-1]
    return f"{int(rev,2):02X}"


def bits_list_to_hex(bits_list) -> str:
    return " ".join(reverse_bits_to_hex(b) for b in bits_list)


def normalize_hex_no_spaces(s: str) -> str:
    return (s or "").replace(" ", "").upper()


def _extract_component_from_line(line: str):
    anchor = line.find(" -- ")
    if anchor != -1:
        after = line[anchor + len(" -- ") :].strip()
    else:
        m = re.match(
            r"^(?:\w{3}\s+)?"               # optional weekday
            r"\d{2}-\d{2}(?:-\d{4})?\s+"    # MM-DD or MM-DD-YYYY
            r"\d{2}:\d{2}:\d{2}\.\d{2}\s+"  # HH:MM:SS.ff
            r"(?P<rest>.+)$",
            line,
        )
        if not m:
            return None
        after = m.group("rest").strip()
    if not after:
        return None
    head = after[:16].lower()
    if head.startswith("ind(") or head.startswith("ctl(") or head.startswith("execute"):
        return None
    return after


def _flush_ixl_group(acc_label, acc_bits, acc_time, out_msgs, start_dt=None, end_dt=None, acc_components=None):
    if not acc_label or not acc_bits:
        return None
    ws_type = IXL_LABEL_TO_WS_TYPE.get(acc_label.upper())
    if not ws_type:
        return None
    include = True
    if start_dt or end_dt:
        try:
            t = parse_time_flexible(acc_time)
            if start_dt and t.time() < start_dt.time():
                include = False
            if end_dt and t.time() > end_dt.time():
                include = False
        except Exception:
            include = False
    if not include:
        return None
    data_hex = bits_list_to_hex(acc_bits)
    comp_lines = acc_components or []
    out_msgs.append({
        "time_tag": acc_time or "",
        "msg_type": ws_type,
        "data_hex": data_hex,
        "_data_norm": normalize_hex_no_spaces(data_hex),
        "component_lines": comp_lines,
    })
    return len(out_msgs) - 1


def parse_ixl_file(ixl_file_path: str, start_dt=None, end_dt=None):
    if not ixl_file_path or not os.path.isfile(ixl_file_path):
        return []

    def _within_point_one_seconds(t1_str: str, t2_str: str) -> bool:
        try:
            t1 = parse_time_flexible(t1_str)
            t2 = parse_time_flexible(t2_str)
            return abs((t2 - t1).total_seconds()) <= 0.1
        except Exception:
            return False

    def _time_from_any_line(s: str):
        m = TIME_PATTERN.search(s)
        return m.group(0) if m else None

    messages = []
    acc_label = None
    acc_bits = []
    acc_time = None
    acc_components = []
    awaiting_components = False
    last_msg_idx = None

    with open(ixl_file_path, "r", encoding="utf-8", errors="ignore") as f:
        lines = f.readlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip("\n")
        m = IXL_LINE_RE.search(line)
        if m:
            awaiting_components = False
            last_msg_idx = None
            label = m.group("label").upper()
            bits = m.group("bits")
            ttag = m.group("hms")

            if acc_label is None:
                acc_label = label
                acc_time = ttag
                acc_bits.append(bits)
                acc_components = []
                j = i
                lookahead_count = 0
                while lookahead_count < 2 and (j + 1) < n:
                    nxt = lines[j + 1].rstrip("\n")
                    nm = IXL_LINE_RE.search(nxt)
                    if nm:
                        nxt_label = nm.group("label").upper()
                        nxt_time = nm.group("hms")
                        if nxt_label == label and _within_point_one_seconds(ttag, nxt_time):
                            acc_bits.append(nm.group("bits"))
                            j += 1
                            lookahead_count += 1
                            continue
                    comp = _extract_component_from_line(nxt)
                    if comp:
                        comp_time = _time_from_any_line(nxt)
                        if comp_time and _within_point_one_seconds(ttag, comp_time):
                            acc_components.append(comp)
                            j += 1
                            continue
                    break
                i = j + 1
                continue
            else:
                if label == acc_label and _within_point_one_seconds(acc_time, ttag):
                    j = i
                    acc_bits.append(bits)
                    lookahead_count = 0
                    while lookahead_count < 2 and (j + 1) < n:
                        nxt = lines[j + 1].rstrip("\n")
                        nm = IXL_LINE_RE.search(nxt)
                        if nm:
                            nxt_label = nm.group("label").upper()
                            nxt_time = nm.group("hms")
                            if nxt_label == acc_label and _within_point_one_seconds(acc_time, nxt_time):
                                acc_bits.append(nm.group("bits"))
                                j += 1
                                lookahead_count += 1
                                continue
                        comp = _extract_component_from_line(nxt)
                        if comp:
                            comp_time = _time_from_any_line(nxt)
                            if comp_time and _within_point_one_seconds(acc_time, comp_time):
                                acc_components.append(comp)
                                j += 1
                                continue
                        break
                    i = j + 1
                    continue
                else:
                    idx = _flush_ixl_group(acc_label, acc_bits, acc_time, messages, start_dt, end_dt, acc_components)
                    if idx is not None and (messages[idx]["msg_type"] in ("Control (12 01)", "Ind (12 8B)")):
                        awaiting_components = True
                        last_msg_idx = idx
                    else:
                        awaiting_components = False
                        last_msg_idx = None

                    acc_label, acc_bits, acc_time = label, [], ttag
                    acc_components = []
                    j = i
                    acc_bits.append(bits)
                    lookahead_count = 0
                    while lookahead_count < 2 and (j + 1) < n:
                        nxt = lines[j + 1].rstrip("\n")
                        nm = IXL_LINE_RE.search(nxt)
                        if nm:
                            nxt_label = nm.group("label").upper()
                            nxt_time = nm.group("hms")
                            if nxt_label == acc_label and _within_point_one_seconds(acc_time, nxt_time):
                                acc_bits.append(nm.group("bits"))
                                j += 1
                                lookahead_count += 1
                                continue
                        comp = _extract_component_from_line(nxt)
                        if comp:
                            comp_time = _time_from_any_line(nxt)
                            if comp_time and _within_point_one_seconds(acc_time, comp_time):
                                acc_components.append(comp)
                                j += 1
                                continue
                        break
                    i = j + 1
                    continue

        if IXL_END_EXECUTE_RE.search(line):
            idx = _flush_ixl_group(acc_label, acc_bits, acc_time, messages, start_dt, end_dt, acc_components)
            if idx is not None and (messages[idx]["msg_type"] in ("Control (12 01)", "Ind (12 8B)")):
                awaiting_components = True
                last_msg_idx = idx
            else:
                awaiting_components = False
                last_msg_idx = None
            acc_label, acc_bits, acc_time = None, [], None
            acc_components = []
            i += 1
            continue

        if acc_bits:
            comp = _extract_component_from_line(line)
            if comp:
                m = TIME_PATTERN.search(line)
                comp_time = m.group(0) if m else None
                if comp_time and acc_time:
                    try:
                        if abs((parse_time_flexible(comp_time) - parse_time_flexible(acc_time)).total_seconds()) <= 0.1:
                            acc_components.append(comp)
                            i += 1
                            continue
                    except Exception:
                        pass
            idx = _flush_ixl_group(acc_label, acc_bits, acc_time, messages, start_dt, end_dt, acc_components)
            if idx is not None and (messages[idx]["msg_type"] in ("Control (12 01)", "Ind (12 8B)")):
                awaiting_components = True
                last_msg_idx = idx
            else:
                awaiting_components = False
                last_msg_idx = None
            acc_label, acc_bits, acc_time = None, [], None
            acc_components = []

        if awaiting_components and last_msg_idx is not None:
            comp = _extract_component_from_line(line)
            if comp:
                messages[last_msg_idx]["component_lines"].append(comp)
                i += 1
                continue
        i += 1

    if acc_bits:
        _flush_ixl_group(acc_label, acc_bits, acc_time, messages, start_dt, end_dt, acc_components)

    for msg in messages:
        lines_join = msg.get("component_lines", [])
        msg["component"] = "\r\n".join(lines_join) if lines_join else ""
        msg.pop("component_lines", None)
    return messages

# ========================
# Excel helpers
# ========================

def adjust_ixl_component_column(
    ws,
    header_row_idx: int = 2,
    start_row: int = 2,
    ixl_component_col_index: int = 6,
    min_width: int = 28,
    max_width: int = 60,
):
    col_letter = get_column_letter(ixl_component_col_index)
    for row in ws.iter_rows(min_row=start_row, max_row=ws.max_row, min_col=ixl_component_col_index, max_col=ixl_component_col_index):
        for cell in row:
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    longest_line_len = 0
    for row in ws.iter_rows(min_row=start_row, max_row=ws.max_row, min_col=ixl_component_col_index, max_col=ixl_component_col_index):
        for cell in row:
            val = str(cell.value) if cell.value is not None else ""
            for ln in val.splitlines():
                longest_line_len = max(longest_line_len, len(ln))
    target_width = max(min_width, min(max_width, longest_line_len + 2))
    ws.column_dimensions[col_letter].width = target_width
    fixed_height = 15
    for r in range(start_row, ws.max_row + 1):
        ws.row_dimensions[r].height = fixed_height

# ========================
# IXL DF aligned to WS
# ========================

def build_ixl_dataframe(ixl_file_path: str, ws_entries: list, start_time_str: str = None, end_time_str: str = None, max_diff_minutes: float = MAX_IXL_WS_DIFF_MINUTES) -> pd.DataFrame:
    row_count = len(ws_entries)
    aligned_empty = [
        {"time tag": "", "msg type": "", "data (hex)": "", "direction": "", "component": ""}
        for _ in range(row_count)
    ]
    if not ixl_file_path or not os.path.isfile(ixl_file_path):
        return pd.DataFrame(aligned_empty, columns=["time tag", "msg type", "data (hex)", "direction", "component"])

    start_dt = end_dt = None
    try:
        if start_time_str:
            start_dt = parse_time_flexible(start_time_str)
        if end_time_str:
            end_dt = parse_time_flexible(end_time_str)
    except Exception:
        start_dt = end_dt = None

    ixl_msgs = parse_ixl_file(ixl_file_path, start_dt=start_dt, end_dt=end_dt)
    unmatched_ixl = [
        {
            "time_tag": m["time_tag"],
            "msg_type": m["msg_type"],
            "component": m.get("component", ""),
            "data_pretty": m["data_hex"],
            "data_norm": normalize_hex_no_spaces(m["data_hex"]),
        }
        for m in ixl_msgs
    ]

    rows = [
        {"time tag": "", "msg type": "", "data (hex)": "", "direction": "", "component": ""}
        for _ in range(row_count)
    ]

    for i, ws_row in enumerate(ws_entries):
        if not ws_row or len(ws_row) < 4:
            continue
        ws_time_str = ws_row[0] or ""
        ws_msg_type = ws_row[1]
        ws_payload_pretty = ws_row[3]
        if not ws_msg_type or not ws_payload_pretty or not ws_time_str:
            continue
        if ws_msg_type not in ("Control (12 01)", "Ind (12 8B)"):
            continue
        ws_payload_norm = normalize_hex_no_spaces(ws_payload_pretty)
        try:
            ws_dt = parse_time_flexible(ws_time_str)
        except Exception:
            continue
        found_idx = None
        for j, im in enumerate(unmatched_ixl):
            if im["msg_type"] != ws_msg_type:
                continue
            if im["data_norm"] != ws_payload_norm:
                continue
            try:
                ixl_dt = parse_time_flexible(im["time_tag"])
            except Exception:
                continue
            if minutes_abs(ws_dt, ixl_dt) <= float(max_diff_minutes):
                found_idx = j
                break
        if found_idx is not None:
            im = unmatched_ixl.pop(found_idx)
            rows[i]["time tag"] = im["time_tag"]
            rows[i]["msg type"] = im["msg_type"]
            rows[i]["data (hex)"] = im["data_pretty"]
            component_text = im.get("component", "")
            if component_text:
                component_text = "See more ...\r\n" + component_text
            rows[i]["component"] = component_text

    return pd.DataFrame(rows, columns=["time tag", "msg type", "data (hex)", "direction", "component"])

# ========================
# Address discovery + formatters
# ========================

def collect_atcs_addresses_starting_with7(pcap_path: str) -> list:
    if not os.path.isfile(pcap_path):
        return []

    cmd = [
        "tshark", "-r", pcap_path,
        "-Y", "atcsl3.dest_addr || atcsl3.srce_addr",
        "-T", "fields",
        "-E", "separator=,",
        "-n", "-q",
        "-e", "atcsl3.dest_addr",
        "-e", "atcsl3.srce_addr",
    ]

    try:
        proc = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            check=True,
        )
    except FileNotFoundError:
        return []
    except subprocess.CalledProcessError:
        return []

    addrs = set()
    for line in proc.stdout.splitlines():
        for token in [t.strip() for t in line.split(",")]:
            if token and token.lower().startswith("7"):
                addrs.add(token)

    return sorted(addrs)


def to_dotted_atcs_format(addr_str: str) -> str:
    hex_only = re.sub(r'[^0-9A-Fa-f]', '', addr_str).lower()
    if len(hex_only) < 10:
        return ""
    hex10 = hex_only[:10]
    hex10 = hex10.replace('a', '0')
    return f"{hex10[0]}.{hex10[1:4]}.{hex10[4:7]}.{hex10[7:10]}"

# ========================
# PCAP time bounds helper
# ========================

def _format_hhmmssmmm(dt: datetime) -> str:
    return dt.strftime("%H:%M:%S.%f")[:-3]


def _parse_hhmmssmmm(s: str):
    try:
        if "." in s:
            hms, frac = s.split(".", 1)
            if len(frac) == 1:
                s = f"{hms}.{frac}00"
            elif len(frac) == 2:
                s = f"{hms}.{frac}0"
            else:
                s = f"{hms}.{frac[:3]}"
        else:
            s = s + ".000"
        return datetime.strptime(s, "%H:%M:%S.%f")
    except Exception:
        return None


def find_pcap_time_bounds(file_path: str):
    if not file_path or not os.path.isfile(file_path):
        return (None, None)
    try:
        packets = rdpcap(file_path)
    except Exception:
        return (None, None)

    tmin = tmax = None
    for pkt in packets:
        try:
            dt = datetime.fromtimestamp(float(pkt.time))
            t_str = dt.strftime("%H:%M:%S.%f")[:-3]
            dt_only = _parse_hhmmssmmm(t_str)
            if not dt_only:
                continue
            if tmin is None or dt_only < tmin:
                tmin = dt_only
            if tmax is None or dt_only > tmax:
                tmax = dt_only
        except Exception:
            continue
    return (_format_hhmmssmmm(tmin) if tmin else None, _format_hhmmssmmm(tmax) if tmax else None)

# ========================
# Core analysis (unchanged behavior)
# ========================

def analyze_logs(ixl_file_path, log_file_path, pcap_file_path, ixl_excel_file_path, start_time_str,
    end_time_str, packetswitch_file_path, target_address, filename_suffix):

    output_file_path = (f"log_packet_analysis_output_{filename_suffix or datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")

    if os.path.exists(output_file_path) and is_file_open(output_file_path):
        messagebox.showerror("Error", "The output file is currently open. Please close it and try again.")
        return

    cm_entries = []
    ws_entries = []
    time_differences = []
    last_two_sequences = []
    recent_sequence_times = {}

    component_map = {}

    if packetswitch_file_path and ixl_excel_file_path:
        try:
            component_map = build_component_map(
                packetswitch_file_path,
                ixl_excel_file_path
            )
            print(f"Component map size: {len(component_map)}")
        except Exception as e:
            print("BitSwitchingDetector failed:", e)


    if log_file_path != False:
        try:
            start_dt = parse_time_only(start_time_str)
            end_dt = parse_time_only(end_time_str)
        except Exception as e:
            messagebox.showerror("Error", f"Invalid time format: {e}")
            return

        with open(log_file_path, "r") as log_file:
            lines = log_file.readlines()
        
        packets = None
        if pcap_file_path:
            packets = rdpcap(pcap_file_path)


        html_text = ""
        if packetswitch_file_path and packetswitch_file_path.lower().endswith(".html"):
            with open(packetswitch_file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                html_text = soup.get_text()
        elif packetswitch_file_path and packetswitch_file_path.lower().endswith(".docx"):
            if is_file_open(packetswitch_file_path):
                messagebox.showerror("Error", "The Packetswitch file is currently open. Please close it and try again.")
                return
            try:
                html_text = "\n".join([para.text for para in Document(packetswitch_file_path).paragraphs])
            except Exception:
                with open(packetswitch_file_path, "r", encoding="utf-8", errors="ignore") as f:
                    html_text = f.read()

        i = 0
        while i < len(lines):
            line = lines[i]
            time_match = TIME_PATTERN.search(line)
            if time_match:
                last_timestamp = time_match.group()
                try:
                    log_dt = parse_time_only(last_timestamp)
                except Exception:
                    i += 1
                    continue
                if not (start_dt <= log_dt <= end_dt):
                    i += 1
                    continue
            else:
                i += 1
                continue

            if "02 02" in line:
                try:
                    index = line.index("02 02")
                    after = line[index + len("02 02") :].strip()
                    after_parts = after.split()
                    if len(after_parts) >= 3 and after_parts[0] == "02":
                        msg_number = "02"
                        msg_type_raw = f"{after_parts[1]} {after_parts[2]}"
                    else:
                        pre_match = re.search(r"(\S{2})\s+02\s+02", line)
                        post_match = re.search(r"02\s+02\s+(\S{2})\s+(\S{2})", line)
                        msg_number = pre_match.group(1) if pre_match else None
                        msg_type_raw = (
                            f"{post_match.group(1)} {post_match.group(2)}" if post_match else None
                        )
                    line_type = "BASIC" if "BASIC" in line else "INFO" if "INFO" in line else ""
                    msg_type = msg_type_raw
                    if msg_type_raw == "12 8B":
                        msg_type = "Ind (12 8B)"
                    elif msg_type_raw == "12 01":
                        msg_type = "Control (12 01)"
                    elif msg_type_raw == "12 48":
                        msg_type = "Recall (12 48)"
                    hex_sequence = f"{msg_number} 02 02 {msg_type_raw}" if msg_number and msg_type_raw else None
                except Exception:
                    i += 1
                    continue

                if not hex_sequence or hex_sequence in last_two_sequences:
                    i += 1
                    continue
                if hex_sequence in recent_sequence_times:
                    if (log_dt - recent_sequence_times[hex_sequence]) <= timedelta(seconds=30):
                        i += 1
                        continue
                try:
                    search_bytes = bytes.fromhex(hex_sequence)
                except Exception:
                    i += 1
                    continue

                if packets:
                    pcap_ts_str, time_diff_str, _, ws_hex_data = find_pcap_time(
                        search_bytes, log_dt, packets, target_address, pcap_file_path, msg_type_raw)
                else:
                    pcap_ts_str = None
                    time_diff_str = ""
                    ws_hex_data = ""


                if time_diff_str != "Found but overflow":
                    hex_data = extract_hex_data(lines, i, msg_type_raw, line_type)
                    cm_entries.append([last_timestamp, msg_type, msg_number, hex_data, ""])
                    
                    if packets:
                        ws_entries.append([
                            pcap_ts_str or "",
                            "" if pcap_ts_str is None else msg_type,
                            "" if pcap_ts_str is None else msg_number,
                            ws_hex_data,
                            ""
                        ])
                    else:
                        # No Wireshark → leave columns empty
                        ws_entries.append(["", "", "", "", ""])

                    time_differences.append([time_diff_str])

                    last_two_sequences.append(hex_sequence)
                    if len(last_two_sequences) > 2:
                        last_two_sequences.pop(0)
                    recent_sequence_times[hex_sequence] = log_dt
                i += 1
            else:
                i += 1

    last_rf_ack_time = None
    if log_file_path == False:
        found_cm_times = set()
    else:
        found_cm_times = set(entry[0] for entry in cm_entries)

    additional_ws_entries = []
    additional_cm_entries = []
    additional_time_differences = []

    
    if log_file_path == False and pcap_file_path:
        packets = rdpcap(pcap_file_path)

        try:
            start_dt = parse_time_only(start_time_str)
            end_dt = parse_time_only(end_time_str)
        except Exception as e:
            messagebox.showerror("Error", f"Invalid time format: {e}")
            return
        filtered_packets = []
        for pkt in packets:
            try:
                pkt_time = datetime.fromtimestamp(float(pkt.time))
                pkt_time_only = datetime.strptime(pkt_time.strftime("%H:%M:%S.%f")[:-3], "%H:%M:%S.%f")
                if start_dt <= pkt_time_only <= end_dt:
                    filtered_packets.append(pkt)
            except Exception:
                continue
        packets = filtered_packets

    html_text = ""
    if packetswitch_file_path and os.path.isfile(packetswitch_file_path):
        if packetswitch_file_path.lower().endswith(".html"):
            with open(packetswitch_file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                html_text = soup.get_text()
        elif packetswitch_file_path.lower().endswith(".docx"):
            if is_file_open(packetswitch_file_path):
                messagebox.showerror("Error", "The Packetswitch file is currently open. Please close it and try again.")
                return
            try:
                html_text = "\n".join([para.text for para in Document(packetswitch_file_path).paragraphs])
            except Exception:
                with open(packetswitch_file_path, "r", encoding="utf-8", errors="ignore") as f:
                    html_text = f.read()

    if pcap_file_path:
        for pkt in rdpcap(pcap_file_path):
            if Raw in pkt:
                data_bytes = pkt[Raw].load
                if is_valid_packet(pkt, data_bytes, target_address):
                    hex_str = data_bytes.hex().upper()
                    hex_parts = [hex_str[i:i+2] for i in range(0, len(hex_str), 2)]
                    for i in range(len(hex_parts) - 3):
                        for j in range(len(hex_parts) - 4):
                            if hex_parts[j] == "8C":
                                fourth_pair = hex_parts[j + 4]
                                if fourth_pair in ["34", "38"]:
                                    rf_ack_time = datetime.fromtimestamp(float(pkt.time))
                                    if last_rf_ack_time and (rf_ack_time - last_rf_ack_time) <= timedelta(seconds=5):
                                        continue
                                    msg_type = "RF_ACK (Inbound)" if fourth_pair == "34" else "RF_ACK (Outbound)"
                                    last_rf_ack_time = rf_ack_time
                                    pcap_time = rf_ack_time.strftime("%H:%M:%S.%f")[:-3]
                                    additional_cm_entries.append(["Not found", "", "", "", ""])
                                    additional_ws_entries.append([pcap_time, msg_type, "", "", ""])
                                    additional_time_differences.append([""])
                                    break
                        if hex_parts[i] == "02" and hex_parts[i + 1] == "02":
                            suffix = target_address.hex()[-2:]
                            search_window = hex_parts[max(0, i - 11) : i]
                            # if suffix not in search_window:
                                # continue
                            if i > 0 and hex_parts[i + 2] == "02":
                                msg_number = "02"
                                msg_type_raw = f"{hex_parts[i + 3]} {hex_parts[i + 4]}"
                            else:
                                msg_number = hex_parts[i - 1] if i > 0 else None
                                msg_type_raw = f"{hex_parts[i + 2]} {hex_parts[i + 3]}"
                            msg_type = msg_type_raw
                            if msg_type_raw == "04 D0":
                                continue
                            if msg_type_raw == "12 8B":
                                msg_type = "Ind (12 8B)"
                            elif msg_type_raw == "12 01":
                                msg_type = "Control (12 01)"
                            elif msg_type_raw == "12 48":
                                msg_type = "Recall (12 48)"
                            hex_sequence = f"{msg_number} 02 02 {msg_type_raw}" if msg_number and msg_type_raw else None

                            if log_file_path != False:
                                if hex_sequence in [
                                    f"{entry[2]} 02 02 {entry[1].split('(')[-1].strip(')')}" for entry in cm_entries if entry[0] != "Not found"
                                ]:
                                    continue
                            pcap_time = datetime.fromtimestamp(float(pkt.time)).strftime("%H:%M:%S.%f")[:-3]

                            if not hex_sequence or hex_sequence in last_two_sequences:
                                continue
                            if hex_sequence in recent_sequence_times:
                                if (datetime.fromtimestamp(float(pkt.time)) - recent_sequence_times[hex_sequence]) <= timedelta(seconds=30):
                                    continue
                            last_two_sequences.append(hex_sequence)
                            if len(last_two_sequences) > 2:
                                last_two_sequences.pop(0)
                            recent_sequence_times[hex_sequence] = datetime.fromtimestamp(float(pkt.time))
                            if pcap_time not in found_cm_times:
                                additional_cm_entries.append(["Not found", "", "", "", ""])
                            fallback_hex = " ".join([data_bytes.hex()[k:k+2] for k in range(0, len(data_bytes.hex()), 2)])
                            ws_hex_data = extract_ws_hex_data(pcap_file_path, pcap_time, msg_type_raw, fallback_hex)
                            additional_ws_entries.append([pcap_time, msg_type, msg_number, ws_hex_data, ""])
                            additional_time_differences.append([""])

    cm_entries.extend(additional_cm_entries)
    ws_entries.extend(additional_ws_entries)
    time_differences.extend(additional_time_differences)

    if log_file_path != False:
        filtered_entries = []
        for i, entry in enumerate(cm_entries):
            cm_time_str = entry[0]
            ws_time_str = ws_entries[i][0]
            try:
                cm_time = parse_time_only(cm_time_str) if cm_time_str != "Not found" else None
                ws_time = parse_time_only(ws_time_str) if ws_time_str else None
                if (cm_time and start_dt <= cm_time <= end_dt) or (ws_time and start_dt <= ws_time <= end_dt):
                    filtered_entries.append((entry, ws_entries[i], time_differences[i]))
            except Exception:
                continue
        cm_entries = [e[0] for e in filtered_entries]
        ws_entries = [e[1] for e in filtered_entries]
        time_differences = [e[2] for e in filtered_entries]

    combined_entries = list(zip(cm_entries, ws_entries, time_differences))

    valid_entries = [e for e in combined_entries if e[0][0] != "Not found"]
    valid_entries.sort(key=lambda x: parse_time_only(x[0][0]))

    did_not_find_entries = [e for e in combined_entries if e[0][0] == "Not found"]
    for entry in did_not_find_entries:
        ws_time_str = entry[1][0]
        try:
            ws_time = parse_time_only(ws_time_str)
        except Exception:
            ws_time = datetime.max
        inserted = False
        for i, valid_entry in enumerate(valid_entries):
            try:
                next_ws_time = parse_time_only(valid_entry[1][0])
                if ws_time < next_ws_time:
                    valid_entries.insert(i, entry)
                    inserted = True
                    break
            except Exception:
                continue
        if not inserted:
            valid_entries.append(entry)

    cm_entries = [e[0] for e in valid_entries]
    ws_entries = [e[1] for e in valid_entries]
    time_differences = [e[2] for e in valid_entries]

    for entry in ws_entries:
        msg_type_raw = entry[1].split("(")[-1].strip(")") if "(" in entry[1] else entry[1]
        entry[3] = process_ws_hex_data(entry[3], msg_type_raw)

    row_count = len(cm_entries)
    ixl_df = build_ixl_dataframe(ixl_file_path, ws_entries, start_time_str, end_time_str)

    diff2_df = pd.DataFrame({" ": [""] * row_count})
    cm_df = pd.DataFrame(cm_entries, columns=["time tag", "msg type", "msg number", "data (hex)", "direction"])
    ws_df = pd.DataFrame(ws_entries, columns=["time tag", "msg type", "msg number", "data (hex)", "UDP/TCP"])
    diff_df = pd.DataFrame(time_differences, columns=[" "])

    html_text = ""
    if packetswitch_file_path and os.path.isfile(packetswitch_file_path):
        if packetswitch_file_path.lower().endswith(".html"):
            with open(packetswitch_file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                html_text = soup.get_text()
        elif packetswitch_file_path.lower().endswith(".docx"):
            if is_file_open(packetswitch_file_path):
                messagebox.showerror("Error", "The Packetswitch file is currently open. Please close it and try again.")
                return
            try:
                html_text = "\n".join([para.text for para in Document(packetswitch_file_path).paragraphs])
            except Exception:
                with open(packetswitch_file_path, "r", encoding="utf-8", errors="ignore") as f:
                    html_text = f.read()

    if not html_text.strip():
        packetswitch_df = pd.DataFrame({
            "time tag": [""] * len(ws_entries),
            "component": [""] * len(ws_entries),
            "data (hex)": [""] * len(ws_entries),
            "data type": [""] * len(ws_entries),
        })
    else:
        is_generic_report = "Generic Report Results" in html_text
        if not is_generic_report:
            packetswitch_times = []
            packetswitch_codes = []
            packetswitch_count = {}
            last_ws_by_type = {"Ind": None, "Control": None}
            
            # Count WS entries per (time, type)
            ws_time_type_count = {}

            for e in ws_entries:
                t = e[0].split(".")[0] if e[0] else None
                mt = e[1]

                if t:
                    key = (t, mt)
                    ws_time_type_count[key] = ws_time_type_count.get(key, 0) + 1

            for entry in ws_entries:
                msg_type_raw = entry[1].split("(")[-1].strip(")") if "(" in entry[1] else entry[1]
                if msg_type_raw not in ["12 8B", "12 01", "12 48"]:
                    packetswitch_times.append("")
                    packetswitch_codes.append("")
                    continue
                pcap_time = entry[0]
                msg_type = entry[1]
                
                ws_data = entry[3].replace(" ", "").upper()

                if "Ind" in msg_type:
                    key_type = "Ind"
                elif "Control" in msg_type:
                    key_type = "Control"   
                elif "Recall" in msg_type:
                    key_type = "Recall"
                else:
                    key_type = None


                prev_ws = last_ws_by_type.get(key_type)

                time_tag_simple = pcap_time.split(".")[0] if pcap_time else None
                duplicate_count = ws_time_type_count.get((time_tag_simple, msg_type), 1)

                # Recall always allowed
                if key_type == "Recall":
                    is_change = True

                # If ONLY ONE WS entry → always allow
                elif duplicate_count == 1:
                    is_change = True

                # If MULTIPLE entries → require actual change
                else:
                    is_change = (prev_ws is None or ws_data != prev_ws)


                code = ""
                if pcap_time:
                    time_tag = pcap_time.split(".")[0]
                    
                    time_minus_one = (datetime.strptime(time_tag, "%H:%M:%S") - timedelta(seconds=1)).strftime("%H:%M:%S")

                    key = (time_tag, msg_type)
                    count = packetswitch_count.get(key, 0)
                    idx = -1
                    for _ in range(count + 1):
                        idx = html_text.find(time_tag, idx + 1)
                        if idx == -1:
                            break
                    if idx != -1:
                                                
                        valid_match = False

                        while idx != -1:
                            after = html_text[idx:]
                            underscore_idx = after.find("_")

                            if underscore_idx != -1 and len(after) > underscore_idx + 2:
                                raw_code = after[underscore_idx + 1 : underscore_idx + 3]

                                # Check if match is correct type
                                if (raw_code == "CR" and msg_type == "Control (12 01)") or \
                                   (raw_code == "IR" and msg_type == "Ind (12 8B)") or \
                                   (raw_code == "R_" and msg_type == "Recall (12 48)"):

                                    valid_match = True
                                    break  # correct one found

                            # wrong match → look for NEXT occurrence
                            idx = html_text.find(time_tag, idx + 1)

                        #  If no valid match found, invalidate idx
                        if not valid_match:
                            idx = -1
                            
                        after = html_text[idx:]
                        underscore_idx = after.find("_")
                        if underscore_idx != -1 and len(after) > underscore_idx + 2:
                            raw_code = after[underscore_idx + 1 : underscore_idx + 3]
                            if raw_code == "CR" and msg_type == "Control (12 01)":
                                code = "Control"
                            elif raw_code == "IR" and msg_type == "Ind (12 8B)":
                                code = "Ind"
                            elif raw_code == "R_" and msg_type == "Recall (12 48)":
                                code = "Recall"
                            if code:
                                if is_change:
                                    packetswitch_times.append(time_tag)
                                    packetswitch_codes.append(code)
                                    packetswitch_count[key] = count + 1
                                    continue
                                else:
                                    code = ""

                
                # NEW: fallback search for -1 second mismatch
                if not code:
                    key = (time_minus_one, msg_type)
                    count = packetswitch_count.get(key, 0)
                    idx = -1

                    for _ in range(count + 1):
                        idx = html_text.find(time_minus_one, idx + 1)
                        if idx == -1:
                            break

                    if idx != -1:
                        after = html_text[idx:]
                        underscore_idx = after.find("_")

                        if underscore_idx != -1 and len(after) > underscore_idx + 2:
                            raw_code = after[underscore_idx + 1 : underscore_idx + 3]

                            if raw_code == "CR" and msg_type == "Control (12 01)":
                                code = "Control"
                            elif raw_code == "IR" and msg_type == "Ind (12 8B)":
                                code = "Ind"
                            elif raw_code == "R_" and msg_type == "Recall (12 48)":
                                code = "Recall"


                            if code:
                                if is_change:
                                    packetswitch_times.append(time_minus_one)
                                    packetswitch_codes.append(code)
                                    packetswitch_count[key] = count + 1
                                    continue
                                else:
                                    code = ""



                tenth_digit_is_nine = False
                try:
                    fractional = pcap_time.split(".")[-1]
                    if len(fractional) >= 1 and fractional[0] == "9":
                        tenth_digit_is_nine = True
                except Exception:
                    pass
                if not code and tenth_digit_is_nine:
                    adjusted_time = (datetime.strptime(time_tag, "%H:%M:%S") + timedelta(seconds=1)).strftime("%H:%M:%S")
                    key = (adjusted_time, msg_type)
                    count = packetswitch_count.get(key, 0)
                    idx = -1
                    for _ in range(count + 1):
                        idx = html_text.find(adjusted_time, idx + 1)
                        if idx == -1:
                            break
                    if idx != -1:
                        after = html_text[idx:]
                        underscore_idx = after.find("_")
                        if underscore_idx != -1 and len(after) > underscore_idx + 2:
                            raw_code = after[underscore_idx + 1 : underscore_idx + 3]
                            if raw_code == "CR" and msg_type == "Control (12 01)":
                                code = "Control"
                            elif raw_code == "IR" and msg_type == "Ind (12 8B)":
                                code = "Ind"
                            elif raw_code == "R_" and msg_type == "Recall (12 48)":
                                code = "Recall"

                            if code:
                                if is_change:
                                    packetswitch_times.append(adjusted_time)
                                    packetswitch_codes.append(code)
                                    packetswitch_count[key] = count + 1
                                    continue
                                else:
                                    code = ""


                            
                if key_type:
                    last_ws_by_type[key_type] = ws_data

                packetswitch_times.append("")
                packetswitch_codes.append("")


            
            # NEW: post-fix misalignment for duplicate time tags
            for i in range(1, len(ws_entries) - 1):
                curr_time = ws_entries[i][0]
                prev_time = ws_entries[i - 1][0]
                next_time = ws_entries[i + 1][0]

                if not curr_time or not prev_time or not next_time:
                    continue

                curr_simple = curr_time.split(".")[0]
                prev_simple = prev_time.split(".")[0]
                next_simple = next_time.split(".")[0]

                curr_type = ws_entries[i][1]
                prev_type = ws_entries[i - 1][1]
                next_type = ws_entries[i + 1][1]

                # must all be same group (time + type)
                if not (curr_simple == prev_simple == next_simple and curr_type == prev_type == next_type):
                    continue

                # get data
                curr_data = ws_entries[i][3].replace(" ", "").upper()
                prev_data = ws_entries[i - 1][3].replace(" ", "").upper()
                next_data = ws_entries[i + 1][3].replace(" ", "").upper()

                # CONDITION:
                # middle = no change
                # next = change
                if curr_data == prev_data and next_data != curr_data:

                    # middle has PS, next does NOT
                    if packetswitch_codes[i] and not packetswitch_codes[i + 1]:

                        # MOVE the packetswitch info DOWN
                        packetswitch_times[i + 1] = packetswitch_times[i]
                        packetswitch_codes[i + 1] = packetswitch_codes[i]

                        packetswitch_times[i] = ""
                        packetswitch_codes[i] = ""

            components = []

            for i in range(len(packetswitch_times)):
                time_tag = packetswitch_times[i]
                ps_type = packetswitch_codes[i]

                component_val = ""

                if time_tag and ps_type:
                    if ps_type == "Ind":
                        key_type = "Ind"
                    elif ps_type == "Control":
                        key_type = "Control"
                    else:
                        key_type = None

                    if key_type:
                        key = (time_tag, key_type)
                        component_val = component_map.get(key, "")

                components.append(component_val)

            packetswitch_df = pd.DataFrame(
                {
                    "time tag": packetswitch_times,
                    "component": components,
                    "data (hex)": [""] * len(packetswitch_times),
                    "data type": packetswitch_codes,
                }
            )

        else:
            packetswitch_times = []
            packetswitch_components = []
            packetswitch_codes = []
            packetswitch_data = []
            ps_lines = html_text.strip().split("\n")
            parsed_ps_lines = []
            for line in ps_lines:
                match = re.match(r".*?(\d{2}:\d{2}:\d{2})\s+([A-Za-z()]+)\s.*?:\s*([0-9A-Fa-f ]+)$", line)
                if match:
                    time_tag = match.group(1)
                    description = match.group(2)
                    data = match.group(3).strip()
                    parsed_ps_lines.append((time_tag, description, data))
            for ws_entry in ws_entries:
                ws_time_str = ws_entry[0].split(".")[0]
                try:
                    ws_time = datetime.strptime(ws_time_str, "%H:%M:%S")
                except Exception:
                    ws_time = None
                ws_data = ws_entry[3].replace(" ", "").upper()
                found_match = False
                for time_tag, description, ps_data in parsed_ps_lines:
                    ps_data_clean = ps_data.replace(" ", "").upper()
                    if ps_data_clean.startswith(ws_data):
                        try:
                            ps_time = datetime.strptime(time_tag, "%H:%M:%S")
                            time_diff = abs((ps_time - ws_time).total_seconds()) if ws_time else None
                        except Exception:
                            time_diff = None
                        if time_diff is not None and time_diff <= 1:
                            if ws_entry[1] in ("RF_ACK (Outbound)", "RF_ACK (Inbound)"):
                                continue
                            if description == "Indic(RF)":
                                if ws_entry[1] != "Ind (12 8B)":
                                    continue
                                else:
                                    description = "Ind"
                            if description == "Rf":
                                if ws_entry[1] != "Control (12 01)":
                                    continue
                                else:
                                    description = "Control"
                            if description == "Indicate":
                                continue
                            packetswitch_times.append(time_tag)
                            
                            time_key = time_tag  # already HH:MM:SS
                            ps_type = description

                            if ps_type == "Ind":
                                key_type = "Ind"
                            elif ps_type == "Control":
                                key_type = "Control"
                            else:
                                key_type = None

                            component_val = ""

                            if key_type:
                                key = (time_key, key_type)
                                component_val = component_map.get(key, "")

                            packetswitch_components.append(component_val)

                            packetswitch_codes.append(description)
                            packetswitch_data.append(ps_data)
                            found_match = True
                            break
                if not found_match:
                    packetswitch_times.append("")
                    packetswitch_components.append("")
                    packetswitch_codes.append("")
                    packetswitch_data.append("")
            packetswitch_df = pd.DataFrame({
                "time tag": packetswitch_times,
                "component": packetswitch_components,
                "data (hex)": packetswitch_data,
                "data type": packetswitch_codes,
            })

    combined_df = pd.concat([
        pd.DataFrame({"Number": list(range(1, len(cm_df) + 1))}),
        ixl_df,
        diff2_df,
        cm_df,
        diff_df,
        ws_df,
        packetswitch_df,
    ], axis=1)

    header_row = [
        " ",
        "IXL Log",
        "",
        "",
        "",
        "",
        "Time Difference",
        "Communication Manager Log",
        "",
        "",
        "",
        "",
        "Time Difference",
        "Wireshark Log",
        "",
        "",
        "",
        "",
        "Packetswitch Data",
        "",
        "",
        "",
    ]

    with pd.ExcelWriter(output_file_path, engine="openpyxl") as writer:
        pd.DataFrame([header_row]).to_excel(writer, index=False, header=False)
        combined_df.to_excel(writer, index=False, startrow=1)

    wb = load_workbook(output_file_path)
    ws = wb.active

    ws.freeze_panes = "A3"

    ws.merge_cells(start_row=1, start_column=2, end_row=1, end_column=6)
    ws.merge_cells(start_row=1, start_column=7, end_row=1, end_column=7)
    ws.merge_cells(start_row=1, start_column=8, end_row=1, end_column=12)
    ws.merge_cells(start_row=1, start_column=13, end_row=1, end_column=13)
    ws.merge_cells(start_row=1, start_column=14, end_row=1, end_column=18)
    ws.merge_cells(start_row=1, start_column=19, end_row=1, end_column=22)

    for col in [2, 7, 8, 13, 14, 19]:
        cell = ws.cell(row=1, column=col)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=22):
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center")

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=6, max_col=6):
        for cell in row:
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    for col in ws.iter_cols(min_row=2):
        max_length = 0
        column = col[0].column
        column_letter = get_column_letter(column)
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column_letter].width = max_length + 2

    adjust_ixl_component_column(ws, header_row_idx=2, start_row=2, ixl_component_col_index=6, min_width=28, max_width=60)

    yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    orange_fill = PatternFill(start_color="FFA500", end_color="FFA500", fill_type="solid")
    red_fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
    COL_WS_TYPE = 15
    COL_WS_NUM = 16
    last_by_group = {}

    
    # Build lookup: group → list of (row, msg_num, time)
    group_entries = {}

    for r in range(2, ws.max_row + 1):
        ws_type_val = ws.cell(row=r, column=COL_WS_TYPE).value
        ws_num_val = ws.cell(row=r, column=COL_WS_NUM).value
        ws_time_val = ws.cell(row=r, column=14).value  # WS time column

        if not ws_type_val or not ws_num_val or not ws_time_val:
            continue

        group = _group_from_msg_type_cell(ws_type_val)

        try:
            msg_num = int(str(ws_num_val).strip(), 16)
        except Exception:
            continue

        try:
            time_val = parse_time_only(ws_time_val)
        except Exception:
            continue

        group_entries.setdefault(group, []).append((r, msg_num, time_val))


    for r in range(2, ws.max_row + 1):
        ws_type_val = ws.cell(row=r, column=COL_WS_TYPE).value
        ws_num_val = ws.cell(row=r, column=COL_WS_NUM).value
        if not ws_type_val or not ws_num_val:
            continue
        raw_type = ws_type_val.split("(")[-1].strip(")") if "(" in ws_type_val else ws_type_val
        if raw_type.upper() == "08 42":
            continue
        group = _group_from_msg_type_cell(ws_type_val)
        try:
            current = int(str(ws_num_val).strip(), 16)
        except Exception:
            continue
        last = last_by_group.get(group)
        if last is None:
            last_by_group[group] = current
            continue
        expected = (last + 2) % 256

        if current != expected:
            ws.cell(row=r, column=COL_WS_NUM).fill = yellow_fill

            # Check for missing previous message
            prev_expected = (current - 2) % 256

            found_prev = False

            group_list = group_entries.get(group, [])

            for rr, msg_num, time_val in group_list:
                if msg_num == prev_expected:
                    # check time difference
                    if abs((time_val - parse_time_only(ws.cell(row=r, column=14).value)).total_seconds()) <= 90:
                        found_prev = True
                        break

            # If previous NOT found → make RED
            if not found_prev:
                ws.cell(row=r, column=COL_WS_NUM).fill = red_fill

        last_by_group[group] = current


    
    # Track last WS data for Ind and Control
    last_ind_data = None
    last_ctrl_data = None

    # Column indexes (based on your layout)
    COL_WS_TYPE = 15
    COL_WS_DATA = 17

    COL_PS_TIME = 19
    COL_PS_COMPONENT = 20
    COL_PS_DATA = 21
    COL_PS_TYPE = 22

    for r in range(2, ws.max_row + 1):

        ws_type = ws.cell(row=r, column=COL_WS_TYPE).value
        ws_data = ws.cell(row=r, column=COL_WS_DATA).value

        # Skip empty rows
        if not ws_type or not ws_data:
            continue

        # Normalize type
        if "Ind" in str(ws_type):
            current_type = "Ind"
        elif "Control" in str(ws_type):
            current_type = "Control"
        else:
            continue

        # Get last value
        last_val = last_ind_data if current_type == "Ind" else last_ctrl_data

        # Detect change
        changed = False
        if last_val is not None and ws_data != last_val:
            changed = True

        # If changed, check Packetswitch
        if changed:
            ps_time_val = ws.cell(row=r, column=COL_PS_TIME).value

            if not ps_time_val or str(ps_time_val).strip() == "":
                #  Highlight ALL 4 Packetswitch columns
                for col in [COL_PS_TIME, COL_PS_COMPONENT, COL_PS_DATA, COL_PS_TYPE]:
                    ws.cell(row=r, column=col).fill = orange_fill

        # Update last values
        if current_type == "Ind":
            last_ind_data = ws_data
        elif current_type == "Control":
            last_ctrl_data = ws_data


    # Format Packetswitch component column like IXL

    COL_PS_COMPONENT = 20  # Packetswitch component column

    for r in range(2, ws.max_row + 1):
        cell = ws.cell(row=r, column=COL_PS_COMPONENT)
        val = cell.value

        if val and isinstance(val, str):
            # Replace "/" with new lines
            parts = [p.strip() for p in val.split("/") if p.strip()]
            
            if len(parts) > 1:
                formatted_text = "See more ...\n" + "\n".join(parts)
            else:
                formatted_text = parts[0] if parts else ""

            cell.value = formatted_text

            # Wrap and align
            cell.alignment = Alignment(
                wrap_text=True,
                horizontal="center",
                vertical="top"
            )

    # Set fixed width for Packetswitch component column
    ws.column_dimensions[get_column_letter(COL_PS_COMPONENT)].width = 20

    # Set fixed row height for better readability
    for r in range(2, ws.max_row + 1):
        ws.row_dimensions[r].height = 15



    wb.save(output_file_path)
    messagebox.showinfo(
        "Success",
        "Output has been saved. The Excel file is ready to be viewed.",
    )
